from __future__ import annotations

import hashlib
import os
from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePath
from typing import Callable

from .chaptering import detect_chapters
from .models import ImportSecurityReport, ImportSourceResponse


MAX_SOURCE_BYTES = 8_000_000
TEXT_SUFFIXES = {".txt", ".md", ".markdown"}
DOCUMENT_SUFFIXES = {".docx", ".pdf"}
ALLOWED_SUFFIXES = TEXT_SUFFIXES | DOCUMENT_SUFFIXES
ALLOWED_CONTENT_TYPES = {
    ".txt": {"", "text/plain", "application/octet-stream"},
    ".md": {"", "text/markdown", "text/plain", "application/octet-stream"},
    ".markdown": {"", "text/markdown", "text/plain", "application/octet-stream"},
    ".docx": {
        "",
        "application/octet-stream",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
    ".pdf": {"", "application/octet-stream", "application/pdf"},
}


@dataclass(frozen=True)
class ExternalScanResult:
    scanner: str
    verdict: str
    risk_level: str
    checks: list[str]
    warnings: list[str]
    blocked_reasons: list[str]


def _default_external_scan_client(
    *,
    endpoint: str,
    scanner: str,
    filename: str,
    content_type: str,
    sha256: str,
    content: bytes,
) -> ExternalScanResult:
    import httpx

    response = httpx.post(
        endpoint,
        files={"file": (filename, content, content_type or "application/octet-stream")},
        data={"filename": filename, "sha256": sha256, "scanner": scanner},
        timeout=_env_positive_float("IMPORT_SECURITY_EXTERNAL_TIMEOUT_SECONDS", 15.0),
    )
    if response.status_code < 200 or response.status_code >= 300:
        raise RuntimeError(f"external scanner returned HTTP {response.status_code}: {response.text}")
    payload = response.json()
    return ExternalScanResult(
        scanner=str(payload.get("scanner") or scanner),
        verdict=str(payload.get("verdict") or "warning"),
        risk_level=str(payload.get("risk_level") or "medium"),
        checks=[str(item) for item in payload.get("checks") or ["external_scan_completed"]],
        warnings=[str(item) for item in payload.get("warnings") or []],
        blocked_reasons=[str(item) for item in payload.get("blocked_reasons") or []],
    )


EXTERNAL_SCAN_CLIENT: Callable[..., ExternalScanResult] = _default_external_scan_client


def validate_source_file_upload(filename: str, content: bytes) -> str:
    safe_name = PurePath(filename or "source.txt").name
    suffix = PurePath(safe_name).suffix.lower()

    if suffix not in ALLOWED_SUFFIXES:
        raise ValueError("Only .txt, .md, .markdown, .docx, and .pdf source files are supported.")
    if len(content) > MAX_SOURCE_BYTES:
        raise ValueError("Source file is too large. Current local import limit is 8 MB.")
    return safe_name


def scan_source_file_upload(filename: str, content_type: str, content: bytes) -> ImportSecurityReport:
    safe_name = validate_source_file_upload(filename, content)
    suffix = PurePath(safe_name).suffix.lower()
    declared_type = (content_type or "").split(";")[0].strip().lower()
    sha256 = hashlib.sha256(content).hexdigest()
    checks = [
        "extension_allowed",
        "size_within_local_limit",
        "sha256_computed",
        "file_signature_checked",
        "declared_content_type_checked",
    ]
    warnings: list[str] = []
    blocked: list[str] = []
    detected_file_type = _detect_file_type(suffix, content)

    allowed_types = ALLOWED_CONTENT_TYPES.get(suffix, {""})
    if declared_type not in allowed_types:
        warnings.append(f"声明的 MIME 类型 {declared_type or 'empty'} 与扩展名 {suffix} 不匹配。")

    if suffix == ".pdf" and not content.startswith(b"%PDF"):
        blocked.append("PDF 文件头不匹配，拒绝导入。")
    elif suffix == ".docx" and not content.startswith(b"PK"):
        blocked.append("DOCX 文件头不匹配，拒绝导入。")
    elif suffix in TEXT_SUFFIXES:
        if b"\x00" in content[:4096]:
            blocked.append("文本稿件包含二进制空字节，疑似非文本文件。")
        if content.startswith(b"%PDF") or content.startswith(b"PK"):
            blocked.append("文本扩展名与二进制文档签名不一致。")

    external_result = _run_external_import_scan(
        filename=safe_name,
        content_type=declared_type or content_type or "application/octet-stream",
        sha256=sha256,
        content=content,
    )
    if external_result:
        checks.extend(external_result.checks)
        warnings.extend(external_result.warnings)
        blocked.extend(external_result.blocked_reasons)

    verdict = "blocked" if blocked else "warning" if warnings else "clean"
    risk_level = _combined_risk_level(
        "high" if blocked else "medium" if warnings else "low",
        external_result.risk_level if external_result else "low",
    )
    return ImportSecurityReport(
        scanner=_combined_scanner_name(external_result),
        verdict=verdict,
        risk_level=risk_level,
        sha256=sha256,
        extension=suffix,
        detected_file_type=detected_file_type,
        declared_content_type=declared_type,
        checks=checks,
        warnings=warnings,
        blocked_reasons=blocked,
    )


def import_source_file(filename: str, content_type: str, content: bytes) -> ImportSourceResponse:
    safe_name = validate_source_file_upload(filename, content)
    suffix = PurePath(safe_name).suffix.lower()
    warnings: list[str] = []
    security_report = scan_source_file_upload(safe_name, content_type, content)
    if security_report.verdict == "blocked":
        raise ValueError("Source file security scan blocked upload: " + " / ".join(security_report.blocked_reasons))
    warnings.extend(security_report.warnings)

    extracted = _extract_source_text(safe_name, suffix, content)
    text = _strip_markdown_document_title(_normalize_source_text(extracted.text), suffix)
    if not text.strip():
        raise ValueError("Source file is empty after decoding.")
    warnings.extend(extracted.warnings)

    chapters = detect_chapters(text)
    paragraph_count = sum(len(chapter.paragraphs) for chapter in chapters)
    if len(chapters) < 3:
        warnings.append("导入文本少于 3 个章节，不满足赛题完整输入要求。")
    if paragraph_count < len(chapters):
        warnings.append("章节段落较少，后续证据索引可能偏弱。")

    return ImportSourceResponse(
        title=_title_from_filename(safe_name),
        text=text,
        filename=safe_name,
        content_type=content_type or "text/plain",
        size_bytes=len(content),
        sha256=security_report.sha256,
        detected_encoding=extracted.detected_encoding,
        extraction_method=extracted.extraction_method,
        document_stats=extracted.document_stats,
        chapter_count=len(chapters),
        paragraph_count=paragraph_count,
        warnings=warnings,
        security_report=security_report,
    )


class ExtractedSource:
    def __init__(
        self,
        text: str,
        detected_encoding: str,
        extraction_method: str,
        document_stats: dict[str, int | str] | None = None,
        warnings: list[str] | None = None,
    ) -> None:
        self.text = text
        self.detected_encoding = detected_encoding
        self.extraction_method = extraction_method
        self.document_stats = document_stats or {}
        self.warnings = warnings or []


def _extract_source_text(filename: str, suffix: str, content: bytes) -> ExtractedSource:
    if suffix in TEXT_SUFFIXES:
        text, encoding = _decode_text(content)
        return ExtractedSource(
            text=text,
            detected_encoding=encoding,
            extraction_method="plain-text" if suffix == ".txt" else "markdown-text",
            document_stats={"characters": len(text)},
        )
    if suffix == ".docx":
        return _extract_docx(content)
    if suffix == ".pdf":
        return _extract_pdf(content)
    raise ValueError(f"Unsupported source file extension: {suffix}")


def _detect_file_type(suffix: str, content: bytes) -> str:
    if content.startswith(b"%PDF"):
        return "pdf"
    if content.startswith(b"PK"):
        return "zip-docx" if suffix == ".docx" else "zip"
    if suffix in TEXT_SUFFIXES:
        return "markdown-text" if suffix in {".md", ".markdown"} else "plain-text"
    return "unknown"


def _run_external_import_scan(
    *,
    filename: str,
    content_type: str,
    sha256: str,
    content: bytes,
) -> ExternalScanResult | None:
    mode = os.getenv("IMPORT_SECURITY_SCANNER", "local").strip().lower() or "local"
    if mode in {"local", "none", "disabled"}:
        return None
    if mode not in {"external", "clamav", "clamav-http"}:
        raise ValueError(f"Unsupported IMPORT_SECURITY_SCANNER: {mode}")

    scanner_name = os.getenv("IMPORT_SECURITY_EXTERNAL_SCANNER_NAME", "").strip() or (
        "clamav-http" if mode.startswith("clamav") else "external-virus-scan"
    )
    endpoint = os.getenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", "").strip()
    if not endpoint:
        return _external_scan_unavailable(
            scanner=scanner_name,
            reason="IMPORT_SECURITY_EXTERNAL_ENDPOINT is not configured.",
        )

    try:
        result = EXTERNAL_SCAN_CLIENT(
            endpoint=endpoint,
            scanner=scanner_name,
            filename=filename,
            content_type=content_type,
            sha256=sha256,
            content=content,
        )
    except Exception as exc:
        return _external_scan_unavailable(scanner=scanner_name, reason=str(exc))
    return _normalize_external_scan_result(result, scanner_name)


def _external_scan_unavailable(scanner: str, reason: str) -> ExternalScanResult:
    policy = os.getenv("IMPORT_SECURITY_ON_UNAVAILABLE", "block").strip().lower() or "block"
    checks = ["external_scan_unavailable"]
    message = f"外部安全扫描器 {scanner} 不可用：{reason}"
    if policy in {"warn", "warning", "allow"}:
        return ExternalScanResult(
            scanner=scanner,
            verdict="warning",
            risk_level="medium",
            checks=checks,
            warnings=[message],
            blocked_reasons=[],
        )
    return ExternalScanResult(
        scanner=scanner,
        verdict="blocked",
        risk_level="high",
        checks=checks,
        warnings=[],
        blocked_reasons=[message],
    )


def _normalize_external_scan_result(result: ExternalScanResult, fallback_scanner: str) -> ExternalScanResult:
    verdict = result.verdict.strip().lower()
    risk_level = result.risk_level.strip().lower()
    checks = result.checks or ["external_scan_completed"]
    warnings = list(result.warnings)
    blocked_reasons = list(result.blocked_reasons)
    if verdict in {"infected", "malicious", "blocked", "deny"}:
        verdict = "blocked"
    elif verdict in {"suspicious", "warning", "warn"}:
        verdict = "warning"
    elif verdict in {"clean", "ok", "pass", "passed"}:
        verdict = "clean"
    else:
        warnings.append(f"外部安全扫描器返回未知 verdict：{result.verdict}")
        verdict = "warning"
    if risk_level not in {"low", "medium", "high", "critical"}:
        risk_level = "high" if verdict == "blocked" else "medium" if verdict == "warning" else "low"
    if verdict == "blocked" and not blocked_reasons:
        blocked_reasons.append("外部安全扫描器阻断该文件。")
    return ExternalScanResult(
        scanner=result.scanner or fallback_scanner,
        verdict=verdict,
        risk_level=risk_level,
        checks=[str(item) for item in checks],
        warnings=warnings,
        blocked_reasons=blocked_reasons,
    )


def _combined_scanner_name(external_result: ExternalScanResult | None) -> str:
    local = "local-static-import-scan-v1"
    if not external_result:
        return local
    return f"{local}+{external_result.scanner}"


def _combined_risk_level(local_risk: str, external_risk: str) -> str:
    order = {"low": 0, "medium": 1, "high": 2, "critical": 3}
    local_score = order.get(local_risk, 1)
    external_score = order.get(external_risk, 1)
    score = max(local_score, external_score)
    for value, index in order.items():
        if index == score:
            return value
    return "medium"


def _env_positive_float(name: str, default: float) -> float:
    value = os.getenv(name, "").strip()
    if not value:
        return default
    try:
        parsed = float(value)
    except ValueError as exc:
        raise ValueError(f"{name} must be a positive number.") from exc
    if parsed <= 0:
        raise ValueError(f"{name} must be a positive number.")
    return parsed


def _extract_docx(content: bytes) -> ExtractedSource:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX import requires python-docx. Run python -m pip install -r apps/api/requirements.txt.") from exc
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError("DOCX file could not be opened or is corrupted.") from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    table_cells.append(value)
    blocks = paragraphs + table_cells
    return ExtractedSource(
        text="\n\n".join(blocks),
        detected_encoding="docx",
        extraction_method="python-docx",
        document_stats={
            "paragraphs": len(paragraphs),
            "tables": len(document.tables),
            "table_cells": len(table_cells),
        },
    )


def _extract_pdf(content: bytes) -> ExtractedSource:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF import requires pypdf. Run python -m pip install -r apps/api/requirements.txt.") from exc
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise ValueError("PDF file could not be opened or is corrupted.") from exc

    warnings: list[str] = []
    page_texts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = _repair_extracted_text_encoding(page.extract_text() or "").strip()
        except Exception:
            page_text = ""
        if page_text:
            page_texts.append(page_text)
        else:
            warnings.append(f"PDF 第 {index} 页未提取到可用文本，扫描件可能需要 OCR。")
    return ExtractedSource(
        text="\n\n".join(page_texts),
        detected_encoding="pdf",
        extraction_method="pypdf",
        document_stats={
            "pages": len(reader.pages),
            "text_pages": len(page_texts),
        },
        warnings=warnings,
    )


def _repair_extracted_text_encoding(text: str) -> str:
    try:
        repaired = text.encode("latin-1").decode("utf-8")
    except UnicodeError:
        return text
    if _cjk_count(repaired) > _cjk_count(text):
        return repaired
    return text


def _cjk_count(text: str) -> int:
    return sum(1 for char in text if "\u4e00" <= char <= "\u9fff")


def _decode_text(content: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace"), "utf-8-replace"


def _normalize_source_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
    return normalized.strip()


def _strip_markdown_document_title(text: str, suffix: str) -> str:
    if suffix not in {".md", ".markdown"}:
        return text
    lines = text.split("\n")
    if lines and lines[0].startswith("# ") and len(lines[0].strip("# ").strip()) <= 80:
        return "\n".join(lines[1:]).lstrip()
    return text


def _title_from_filename(filename: str) -> str:
    stem = PurePath(filename).stem.strip()
    return stem or "未命名改编项目"
