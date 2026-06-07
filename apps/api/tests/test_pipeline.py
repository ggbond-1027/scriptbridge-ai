import asyncio
from io import BytesIO

from fastapi.testclient import TestClient

from apps.api.chaptering import detect_chapters
from apps.api.evidence import retrieve_scene_evidence
from apps.api.llm_provider import ProviderResult
from apps.api.main import app
from apps.api.pipeline import _build_llm_prompt, build_fallback_screenplay
from apps.api.sample import SAMPLE_NOVEL
from apps.api.models import AdaptationStyle
from apps.api.exporters import to_fountain, to_markdown, to_yaml
from apps.api.validation import validate_screenplay, validate_yaml_text


class FakeRedisClient:
    def __init__(self):
        self.queues: dict[str, list[str]] = {}

    def lpush(self, name: str, *values: str) -> int:
        queue = self.queues.setdefault(name, [])
        for value in values:
            queue.insert(0, value)
        return len(queue)

    def brpop(self, keys, timeout: int = 0):
        name = keys[0] if isinstance(keys, list) else keys
        queue = self.queues.setdefault(name, [])
        if not queue:
            return None
        return name, queue.pop()

    def llen(self, name: str) -> int:
        return len(self.queues.get(name, []))


def _configure_test_storage(monkeypatch, data_dir):
    from apps.api import storage

    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")


def _enable_strict_rbac(monkeypatch) -> None:
    monkeypatch.setenv("AUTH_MODE", "strict-rbac")


def _approved_delivery_project(client: TestClient, title: str = "Delivery package ready"):
    project = client.post("/api/projects", json={"title": title, "source_text": SAMPLE_NOVEL}).json()
    owner_name = next(member["name"] for member in project["members"] if member["role"] == "owner")
    client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": owner_name, "name": "writer_ascii", "role": "writer"},
    )
    client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": owner_name, "name": "producer_ascii", "role": "producer"},
    )
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    version_id = client.get(f"/api/jobs/{generated_job['id']}").json()["result_version_id"]
    submitted = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "writer_ascii", "note": "request delivery approval"},
    ).json()
    approved = client.post(
        f"/api/projects/{project['id']}/approvals/{submitted['approval']['id']}/decision",
        json={"actor": "producer_ascii", "decision": "approve", "note": "approve package"},
    ).json()
    return project, version_id, approved["approval"]


def test_detect_chapters_finds_three_chapters():
    chapters = detect_chapters(SAMPLE_NOVEL)
    assert len(chapters) == 3
    assert chapters[0].id == "ch_001"
    assert chapters[0].paragraphs


def test_fallback_screenplay_is_valid():
    chapters = detect_chapters(SAMPLE_NOVEL)
    screenplay = build_fallback_screenplay(chapters, "雨夜来信", AdaptationStyle())
    result = validate_screenplay(screenplay)
    assert result.valid, result.issues
    assert len(screenplay.scenes) == 3
    assert screenplay.scenes[0].source_refs
    assert screenplay.quality_report.overall_score > 0
    assert screenplay.quality_report.metrics
    assert screenplay.production.estimated_runtime_minutes > 0
    assert screenplay.metadata.pipeline_stages


def test_yaml_roundtrip_validation():
    screenplay = build_fallback_screenplay(detect_chapters(SAMPLE_NOVEL), "雨夜来信", AdaptationStyle())
    yaml_text = to_yaml(screenplay)
    result = validate_yaml_text(yaml_text)
    assert result.valid, result.issues


def test_exporters_return_content():
    screenplay = build_fallback_screenplay(detect_chapters(SAMPLE_NOVEL), "雨夜来信", AdaptationStyle())
    assert "schema_version" in to_yaml(screenplay)
    assert "## 剧本" in to_markdown(screenplay)
    assert "INT." in to_fountain(screenplay) or "EXT." in to_fountain(screenplay)


def test_api_generate_without_key_uses_fallback():
    client = TestClient(app)
    response = client.post("/api/generate", json={"text": SAMPLE_NOVEL, "title": "雨夜来信", "use_llm": False})
    assert response.status_code == 200
    payload = response.json()
    assert payload["validation"]["valid"] is True
    assert payload["screenplay"]["metadata"]["source_chapter_count"] == 3
    assert payload["screenplay"]["quality_report"]["overall_score"] > 0


def test_api_capabilities_exposes_enterprise_stack():
    client = TestClient(app)
    response = client.get("/api/capabilities")
    assert response.status_code == 200
    payload = response.json()
    assert payload["product_tier"] == "enterprise-v2"
    assert "ChapterIndexer" in payload["pipeline_agents"]
    assert "LocalHashEmbeddings" in payload["pipeline_agents"]
    assert "SystemReadinessGate" in payload["pipeline_agents"]
    assert "FastAPI" in payload["open_source_stack"]


def test_system_readiness_blocks_local_only_defaults(monkeypatch):
    for name in [
        "DATABASE_URL",
        "POSTGRES_DSN",
        "SCRIPTBRIDGE_DATABASE_URL",
        "JOB_BROKER_URL",
        "REDIS_URL",
        "OPENAI_API_KEY",
        "OPENAI_API_KEY_SECRET_REF",
        "MODEL_SECRET_STUDIO_PROD",
        "IMPORT_SECURITY_EXTERNAL_ENDPOINT",
        "DELIVERY_SIGNING_SECRET",
        "AUTH_MODE",
    ]:
        monkeypatch.delenv(name, raising=False)
    monkeypatch.setenv("JOB_QUEUE_MODE", "background")
    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "local")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "local")

    client = TestClient(app)
    response = client.get("/api/system-readiness")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "blocked"
    assert payload["score"] < 60
    check_status = {check["id"]: check["status"] for check in payload["checks"]}
    assert check_status["database"] == "blocked"
    assert check_status["queue"] == "blocked"
    assert check_status["model_secrets"] == "blocked"
    assert check_status["artifact_storage"] == "warning"
    assert payload["next_actions"]
    assert "sk-" not in response.text.lower()


def test_system_readiness_accepts_production_configuration(monkeypatch):
    monkeypatch.setenv("SCRIPTBRIDGE_ENV", "production")
    monkeypatch.setenv("DATABASE_URL", "postgresql://user:pass@db.example.test:5432/scriptbridge")
    monkeypatch.setenv("JOB_QUEUE_MODE", "external")
    monkeypatch.setenv("JOB_BROKER_URL", "redis://redis.example.test:6379/0")
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.setenv("OPENAI_API_KEY_SECRET_REF", "studio-prod")
    monkeypatch.setenv("MODEL_SECRET_STUDIO_PROD", "server-side-secret")
    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "clamav")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", "http://scanner.example.test/scan")
    monkeypatch.setenv("IMPORT_SECURITY_ON_UNAVAILABLE", "block")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "s3")
    monkeypatch.setenv("DELIVERY_ARTIFACT_BUCKET", "delivery-bucket")
    monkeypatch.setenv("DELIVERY_ARTIFACT_ACCESS_KEY_ID", "access-key")
    monkeypatch.setenv("DELIVERY_ARTIFACT_SECRET_ACCESS_KEY", "secret-key")
    monkeypatch.setenv("AUTH_MODE", "oidc")
    monkeypatch.setenv("DELIVERY_SIGNING_SECRET", "delivery-signing-secret")

    client = TestClient(app)
    response = client.get("/api/system-readiness")

    assert response.status_code == 200
    payload = response.json()
    assert payload["environment"] == "production"
    assert payload["status"] == "ready"
    assert payload["score"] == 100
    assert not payload["blockers"]
    check_status = {check["id"]: check["status"] for check in payload["checks"]}
    assert check_status == {
        "database": "ready",
        "queue": "ready",
        "model_secrets": "ready",
        "import_security": "ready",
        "artifact_storage": "ready",
        "auth": "ready",
    }
    queue_check = next(check for check in payload["checks"] if check["id"] == "queue")
    assert queue_check["evidence"]["broker"] == "redis"
    assert queue_check["evidence"]["redis_adapter"] is True
    assert "server-side-secret" not in response.text
    assert "secret-key" not in response.text


def test_import_source_file_detects_chapters_and_metadata():
    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["title"] == "rain-letter"
    assert payload["text"].startswith("第一章")
    assert payload["filename"] == "rain-letter.md"
    assert payload["detected_encoding"] in {"utf-8-sig", "utf-8"}
    assert payload["chapter_count"] == 3
    assert payload["paragraph_count"] >= 9
    assert "# 雨夜来信" not in payload["text"]
    assert payload["sha256"]
    assert payload["security_report"]["verdict"] == "clean"
    assert payload["security_report"]["detected_file_type"] == "markdown-text"


def test_import_source_rejects_unsupported_extension():
    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={"file": ("source.epub", b"PK\x03\x04", "application/epub+zip")},
    )

    assert response.status_code == 422


def test_import_source_security_scan_blocks_mismatched_signature():
    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={"file": ("fake-text.txt", _minimal_pdf_bytes(), "text/plain")},
    )

    assert response.status_code == 422
    assert "security scan blocked" in response.text


def test_import_source_external_security_scan_blocks_infected_file(monkeypatch):
    from apps.api import importer
    from apps.api.importer import ExternalScanResult

    def fake_external_scan(**kwargs):
        assert kwargs["scanner"] == "clamav-http"
        assert kwargs["sha256"]
        return ExternalScanResult(
            scanner="clamav-http",
            verdict="infected",
            risk_level="critical",
            checks=["external_scan_completed"],
            warnings=[],
            blocked_reasons=["Eicar-Test-Signature"],
        )

    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "clamav")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", "http://scanner.example.test/scan")
    monkeypatch.setattr(importer, "EXTERNAL_SCAN_CLIENT", fake_external_scan)

    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 422
    assert "Eicar-Test-Signature" in response.text


def test_import_source_external_security_scan_warning_is_preserved(monkeypatch):
    from apps.api import importer
    from apps.api.importer import ExternalScanResult

    def fake_external_scan(**kwargs):
        return ExternalScanResult(
            scanner="clamav-http",
            verdict="suspicious",
            risk_level="medium",
            checks=["external_scan_completed"],
            warnings=["外部扫描提示低置信度风险。"],
            blocked_reasons=[],
        )

    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "external")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_SCANNER_NAME", "clamav-http")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", "http://scanner.example.test/scan")
    monkeypatch.setattr(importer, "EXTERNAL_SCAN_CLIENT", fake_external_scan)

    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["security_report"]["scanner"] == "local-static-import-scan-v1+clamav-http"
    assert payload["security_report"]["verdict"] == "warning"
    assert payload["security_report"]["risk_level"] == "medium"
    assert "external_scan_completed" in payload["security_report"]["checks"]
    assert "外部扫描提示低置信度风险。" in payload["security_report"]["warnings"]


def test_import_source_external_scanner_unavailable_can_fail_closed(monkeypatch):
    client = TestClient(app)
    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "clamav")
    monkeypatch.delenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", raising=False)

    response = client.post(
        "/api/import/source",
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 422
    assert "不可用" in response.text


def test_import_source_external_scanner_unavailable_can_warn(monkeypatch):
    client = TestClient(app)
    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "clamav")
    monkeypatch.setenv("IMPORT_SECURITY_ON_UNAVAILABLE", "warn")
    monkeypatch.delenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", raising=False)

    response = client.post(
        "/api/import/source",
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["security_report"]["verdict"] == "warning"
    assert "external_scan_unavailable" in payload["security_report"]["checks"]
    assert payload["security_report"]["warnings"]


def test_import_source_docx_extracts_manuscript_text():
    from docx import Document

    document = Document()
    for block in SAMPLE_NOVEL.split("\n\n"):
        document.add_paragraph(block)
    buffer = BytesIO()
    document.save(buffer)

    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={
            "file": (
                "rain-letter.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["title"] == "rain-letter"
    assert payload["detected_encoding"] == "docx"
    assert payload["extraction_method"] == "python-docx"
    assert payload["chapter_count"] == 3
    assert payload["paragraph_count"] >= 9
    assert payload["document_stats"]["paragraphs"] >= 12
    assert "第一章 雨夜来信" in payload["text"]


def test_import_source_pdf_extracts_pages_and_metadata():
    client = TestClient(app)
    response = client.post(
        "/api/import/source",
        files={"file": ("rain-letter.pdf", _minimal_pdf_bytes(), "application/pdf")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["title"] == "rain-letter"
    assert payload["detected_encoding"] == "pdf"
    assert payload["extraction_method"] == "pypdf"
    assert payload["document_stats"]["pages"] == 1
    assert payload["chapter_count"] == 3
    assert "第三章 报社档案室" in payload["text"]


def test_import_source_job_updates_project_source_text(tmp_path, monkeypatch):
    from apps.api import main as main_module
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Import target", "source_text": ""}).json()
    response = client.post(
        "/api/jobs/import-source",
        data={"project_id": project["id"], "actor": "system"},
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 200
    job = response.json()
    assert job["kind"] == "import"
    assert job["status"] == "succeeded"
    assert job["progress"] == 100
    assert job["result_payload"]["filename"] == "rain-letter.md"
    assert job["result_payload"]["chapter_count"] == 3
    assert job["result_payload"]["sha256"] == job["request_payload"]["sha256"]
    assert job["result_payload"]["security_report"]["verdict"] == "clean"
    assert job["request_payload"]["security_report"]["scanner"] == "local-static-import-scan-v1"
    assert "file_path" in job["request_payload"]
    assert "text" not in job["request_payload"]

    detail = client.get(f"/api/projects/{project['id']}").json()
    assert detail["title"] == "rain-letter"
    assert detail["source_text"].startswith("第一章 雨夜来信")
    assert detail["source_text"].rstrip() == SAMPLE_NOVEL.rstrip()
    assert detail["last_job_id"] == job["id"]
    assert detail["audit_events"][0]["event_type"] == "import.source_completed"

    stream = client.get(f"/api/jobs/{job['id']}/events")
    assert stream.status_code == 200
    assert "import.completed" in stream.text


def test_import_source_job_persists_external_security_scan_metadata(tmp_path, monkeypatch):
    from apps.api import importer
    from apps.api import main as main_module
    from apps.api import storage
    from apps.api.importer import ExternalScanResult

    def fake_external_scan(**kwargs):
        return ExternalScanResult(
            scanner="clamav-http",
            verdict="clean",
            risk_level="low",
            checks=["external_scan_completed"],
            warnings=[],
            blocked_reasons=[],
        )

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "external")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_SCANNER_NAME", "clamav-http")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", "http://scanner.example.test/scan")
    monkeypatch.setattr(importer, "EXTERNAL_SCAN_CLIENT", fake_external_scan)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Import external scan", "source_text": ""}).json()
    response = client.post(
        "/api/jobs/import-source",
        data={"project_id": project["id"], "actor": "system"},
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    )

    assert response.status_code == 200
    job = response.json()
    assert job["request_payload"]["security_report"]["scanner"] == "local-static-import-scan-v1+clamav-http"
    assert "external_scan_completed" in job["request_payload"]["security_report"]["checks"]
    assert job["result_payload"]["security_report"]["scanner"] == "local-static-import-scan-v1+clamav-http"

    history = client.get(f"/api/projects/{project['id']}/import-history").json()
    assert history["imports"][0]["security_report"]["scanner"] == "local-static-import-scan-v1+clamav-http"


def test_project_import_history_exposes_sanitized_import_runs(tmp_path, monkeypatch):
    from apps.api import main as main_module
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Import history", "source_text": ""}).json()
    first = client.post(
        "/api/jobs/import-source",
        data={"project_id": project["id"], "actor": "制片"},
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    ).json()
    second = client.post(
        "/api/jobs/import-source",
        data={"project_id": project["id"], "actor": "制片"},
        files={"file": ("rain-letter.txt", SAMPLE_NOVEL.encode("utf-8"), "text/plain")},
    ).json()

    response = client.get(f"/api/projects/{project['id']}/import-history")

    assert response.status_code == 200
    payload = response.json()
    assert payload["project_id"] == project["id"]
    assert payload["total"] == 2
    assert payload["succeeded"] == 2
    assert payload["failed"] == 0
    assert payload["imports"][0]["job_id"] == second["id"]
    assert payload["imports"][0]["filename"] == "rain-letter.txt"
    assert payload["imports"][0]["actor"] == "制片"
    assert payload["imports"][0]["chapter_count"] == 3
    assert payload["imports"][0]["paragraph_count"] == 12
    assert payload["imports"][0]["upload_mode"] == "multipart"
    assert payload["imports"][0]["sha256"]
    assert payload["imports"][0]["security_report"]["verdict"] == "clean"
    assert payload["imports"][0]["last_stage_id"] == "import.completed"
    assert payload["imports"][0]["audit_event_id"]
    assert "file_path" not in response.text
    assert SAMPLE_NOVEL[:24] not in response.text
    assert first["request_payload"]["filename"] == "rain-letter.md"


def test_chunked_import_session_uploads_parts_and_dispatches_import_job(tmp_path, monkeypatch):
    from apps.api import main as main_module
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setattr(main_module, "CHUNK_UPLOAD_DIR", data_dir / "chunked-imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Chunk target", "source_text": ""}).json()
    source = SAMPLE_NOVEL.encode("utf-8")
    part_size = len(source) // 3
    create = client.post(
        "/api/import-sessions",
        json={
            "filename": "chunk-rain.md",
            "content_type": "text/markdown",
            "size_bytes": len(source),
            "total_chunks": 3,
            "project_id": project["id"],
            "actor": "编剧",
        },
    )
    assert create.status_code == 200
    session = create.json()
    assert session["status"] == "pending"
    assert session["uploaded_chunks"] == []

    chunks = [source[:part_size], source[part_size : part_size * 2], source[part_size * 2 :]]
    for index, chunk in enumerate(chunks):
        response = client.put(
            f"/api/import-sessions/{session['id']}/chunks/{index}",
            content=chunk,
            headers={"content-type": "application/octet-stream"},
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["uploaded_count"] == index + 1

    complete = client.post(f"/api/import-sessions/{session['id']}/complete")

    assert complete.status_code == 200
    payload = complete.json()
    assert payload["session"]["status"] == "completed"
    job = payload["job"]
    assert job["kind"] == "import"
    assert job["status"] == "succeeded"
    assert job["request_payload"]["upload_session_id"] == session["id"]
    assert job["request_payload"]["filename"] == "chunk-rain.md"
    assert job["request_payload"]["upload_mode"] == "chunked"
    assert job["request_payload"]["security_report"]["verdict"] == "clean"
    assert "text" not in job["request_payload"]

    detail = client.get(f"/api/projects/{project['id']}").json()
    assert detail["source_text"].rstrip() == SAMPLE_NOVEL.rstrip()
    assert detail["last_job_id"] == job["id"]

    history = client.get(f"/api/projects/{project['id']}/import-history").json()
    assert history["imports"][0]["job_id"] == job["id"]
    assert history["imports"][0]["filename"] == "chunk-rain.md"
    assert history["imports"][0]["upload_mode"] == "chunked"
    assert history["imports"][0]["security_report"]["sha256"] == job["result_payload"]["sha256"]
    assert "chunked-imports" not in client.get(f"/api/projects/{project['id']}/import-history").text


def test_chunked_import_session_external_security_scan_can_block_completion(tmp_path, monkeypatch):
    from apps.api import importer
    from apps.api import main as main_module
    from apps.api import storage
    from apps.api.importer import ExternalScanResult

    def fake_external_scan(**kwargs):
        return ExternalScanResult(
            scanner="clamav-http",
            verdict="blocked",
            risk_level="critical",
            checks=["external_scan_completed"],
            warnings=[],
            blocked_reasons=["Chunked-Eicar-Signature"],
        )

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setattr(main_module, "CHUNK_UPLOAD_DIR", data_dir / "chunked-imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("IMPORT_SECURITY_SCANNER", "external")
    monkeypatch.setenv("IMPORT_SECURITY_EXTERNAL_ENDPOINT", "http://scanner.example.test/scan")
    monkeypatch.setattr(importer, "EXTERNAL_SCAN_CLIENT", fake_external_scan)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Chunk blocked", "source_text": ""}).json()
    source = SAMPLE_NOVEL.encode("utf-8")
    create = client.post(
        "/api/import-sessions",
        json={
            "filename": "chunk-rain.md",
            "content_type": "text/markdown",
            "size_bytes": len(source),
            "total_chunks": 2,
            "project_id": project["id"],
            "actor": "编剧",
        },
    ).json()
    chunks = [source[: len(source) // 2], source[len(source) // 2 :]]
    for index, chunk in enumerate(chunks):
        uploaded = client.put(
            f"/api/import-sessions/{create['id']}/chunks/{index}",
            content=chunk,
            headers={"content-type": "application/octet-stream"},
        )
        assert uploaded.status_code == 200

    complete = client.post(f"/api/import-sessions/{create['id']}/complete")

    assert complete.status_code == 422
    assert "Chunked-Eicar-Signature" in complete.text


def test_import_source_job_can_create_project_and_wait_for_external_worker(tmp_path, monkeypatch):
    from apps.api import main as main_module
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "external")

    client = TestClient(app)
    queued = client.post(
        "/api/jobs/import-source",
        data={"actor": "system"},
        files={"file": ("rain-letter.txt", SAMPLE_NOVEL.encode("utf-8"), "text/plain")},
    )

    assert queued.status_code == 200
    job = queued.json()
    assert job["kind"] == "import"
    assert job["status"] == "queued"
    assert job["queue_mode"] == "external"

    project = client.get(f"/api/projects/{job['project_id']}").json()
    assert project["source_text"] == ""

    worked = client.post("/api/workers/run-once")
    assert worked.status_code == 200
    worked_job = worked.json()["job"]
    assert worked_job["status"] == "succeeded"
    assert worked_job["result_payload"]["extraction_method"] == "plain-text"

    imported_project = client.get(f"/api/projects/{job['project_id']}").json()
    assert imported_project["source_text"].startswith("第一章 雨夜来信")
    assert imported_project["source_text"].rstrip() == SAMPLE_NOVEL.rstrip()
    assert imported_project["last_job_id"] == job["id"]


def test_external_generation_job_uses_configured_redis_broker(tmp_path, monkeypatch):
    from apps.api import redis_broker, storage

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    fake_redis = FakeRedisClient()

    def fake_broker():
        return redis_broker.RedisBroker(
            redis_broker.RedisBrokerConfig(
                url="redis://redis.example.test:6379/0",
                queue_name="scriptbridge:test",
                pop_timeout_seconds=0,
            ),
            client=fake_redis,
        )

    monkeypatch.setenv("JOB_QUEUE_MODE", "external")
    monkeypatch.setenv("JOB_BROKER_URL", "redis://redis.example.test:6379/0")
    monkeypatch.setattr("apps.api.queue.configured_redis_broker", fake_broker)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Redis queue project", "source_text": SAMPLE_NOVEL}).json()
    queued = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})

    assert queued.status_code == 200
    job = queued.json()
    assert job["queue_mode"] == "external"
    assert job["status"] == "queued"
    assert fake_redis.llen("scriptbridge:test") == 1

    status = client.get("/api/job-queue/status").json()
    assert "Redis broker" in status["worker_hint"]
    assert "scriptbridge:test" in status["worker_hint"]

    worked = client.post("/api/workers/run-once")

    assert worked.status_code == 200
    worked_job = worked.json()["job"]
    assert worked_job["id"] == job["id"]
    assert worked_job["status"] == "succeeded"
    assert fake_redis.llen("scriptbridge:test") == 0
    stored = storage.get_job(job["id"])
    assert stored.locked_by == "local-worker"


def test_redis_broker_worker_skips_stale_or_canceled_jobs(tmp_path, monkeypatch):
    from apps.api import redis_broker, storage

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    fake_redis = FakeRedisClient()

    def fake_broker():
        return redis_broker.RedisBroker(
            redis_broker.RedisBrokerConfig(
                url="redis://redis.example.test:6379/0",
                queue_name="scriptbridge:test",
                pop_timeout_seconds=0,
            ),
            client=fake_redis,
        )

    monkeypatch.setenv("JOB_QUEUE_MODE", "external")
    monkeypatch.setenv("JOB_BROKER_URL", "redis://redis.example.test:6379/0")
    monkeypatch.setattr("apps.api.queue.configured_redis_broker", fake_broker)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Redis stale project", "source_text": SAMPLE_NOVEL}).json()
    first = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    second = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    storage.cancel_job(first["id"], "项目负责人")

    worked = client.post("/api/workers/run-once")

    assert worked.status_code == 200
    worked_job = worked.json()["job"]
    assert worked_job["id"] == second["id"]
    assert worked_job["status"] == "succeeded"
    assert storage.get_job(first["id"]).status == "canceled"
    assert fake_redis.llen("scriptbridge:test") == 0


def test_worker_loop_records_heartbeat_and_processes_redis_jobs(tmp_path, monkeypatch):
    import asyncio

    from apps.api import redis_broker, storage
    from apps.api.queue import run_worker_loop

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    fake_redis = FakeRedisClient()

    def fake_broker():
        return redis_broker.RedisBroker(
            redis_broker.RedisBrokerConfig(
                url="redis://redis.example.test:6379/0",
                queue_name="scriptbridge:test",
                pop_timeout_seconds=0,
            ),
            client=fake_redis,
        )

    monkeypatch.setenv("JOB_QUEUE_MODE", "external")
    monkeypatch.setenv("JOB_BROKER_URL", "redis://redis.example.test:6379/0")
    monkeypatch.setattr("apps.api.queue.configured_redis_broker", fake_broker)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Worker loop project", "source_text": SAMPLE_NOVEL}).json()
    queued = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()

    result = asyncio.run(run_worker_loop(worker_id="worker-a", max_jobs=1, idle_sleep_seconds=0, broker=fake_broker()))

    assert result.processed == 1
    assert result.worker.worker_id == "worker-a"
    assert result.worker.completed_jobs == 1
    assert result.worker.status == "stopped"
    assert storage.get_job(queued["id"]).status == "succeeded"

    status = client.get("/api/job-queue/status").json()
    worker = next(item for item in status["workers"] if item["worker_id"] == "worker-a")
    assert worker["completed_jobs"] == 1
    assert worker["failed_jobs"] == 0
    assert worker["status"] == "stopped"


def test_job_exhausting_attempts_moves_to_dead_letter_queue(tmp_path, monkeypatch):
    from apps.api import storage
    from apps.api.queue import execute_job

    async def fail_generation(*args, **kwargs):
        raise RuntimeError("forced generation failure")

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setattr("apps.api.jobs.generate_screenplay", fail_generation)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Dead letter project", "source_text": SAMPLE_NOVEL}).json()
    job = storage.create_job(project["id"], "generate", queue_mode="external", request_payload={"use_llm": False})
    job.max_attempts = 1
    storage.save_job(job)

    completed = asyncio.run(execute_job(job.id))

    assert completed.status == "dead_lettered"
    assert completed.dead_lettered_at
    assert completed.dead_letter_reason
    assert completed.attempts == 1

    status = client.get("/api/job-queue/status").json()
    assert status["dead_lettered"] == 1
    dead_letter = client.get("/api/job-queue/dead-letter").json()
    assert dead_letter["dead_lettered"] == 1
    assert dead_letter["jobs"][0]["id"] == job.id
    assert dead_letter["jobs"][0]["status"] == "dead_lettered"


def test_dead_letter_job_can_requeue_and_execute_through_redis(tmp_path, monkeypatch):
    from apps.api import redis_broker, storage

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    fake_redis = FakeRedisClient()

    def fake_broker():
        return redis_broker.RedisBroker(
            redis_broker.RedisBrokerConfig(
                url="redis://redis.example.test:6379/0",
                queue_name="scriptbridge:test",
                pop_timeout_seconds=0,
            ),
            client=fake_redis,
        )

    monkeypatch.setenv("JOB_QUEUE_MODE", "external")
    monkeypatch.setenv("JOB_BROKER_URL", "redis://redis.example.test:6379/0")
    monkeypatch.setattr("apps.api.queue.configured_redis_broker", fake_broker)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Dead letter retry project", "source_text": SAMPLE_NOVEL}).json()
    owner_name = next(member["name"] for member in project["members"] if member["role"] == "owner")
    client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": owner_name, "name": "ops_owner", "role": "owner"},
    )
    job = storage.create_job(project["id"], "generate", queue_mode="external", request_payload={"use_llm": False})
    job.status = "dead_lettered"
    job.attempts = job.max_attempts
    job.error = "previous failure"
    job.dead_lettered_at = "2026-01-01T00:00:00+00:00"
    job.dead_letter_reason = "previous failure"
    job.dead_letter_source = "worker-a"
    storage.save_job(job)

    retried = client.post(f"/api/jobs/{job.id}", json={"actor": "ops_owner", "action": "retry"})

    assert retried.status_code == 200
    payload = retried.json()
    assert payload["status"] == "queued"
    assert payload["attempts"] == 0
    assert payload["requeue_count"] == 1
    assert payload["dead_lettered_at"] is None
    assert fake_redis.llen("scriptbridge:test") == 1

    worked = client.post("/api/workers/run-once")
    assert worked.status_code == 200
    assert worked.json()["job"]["status"] == "succeeded"
    assert fake_redis.llen("scriptbridge:test") == 0


def test_model_profile_endpoint_exposes_sanitized_defaults(monkeypatch):
    client = TestClient(app)
    monkeypatch.setenv("LLM_PROVIDER", "api")
    monkeypatch.setenv("OPENAI_BASE_URL", "https://api.example.test/v1")
    monkeypatch.setenv("MODEL_NAME", "gpt-5.5")
    monkeypatch.setenv("OPENAI_API_KEY", "secret-profile-key")

    response = client.get("/api/model-profiles")

    assert response.status_code == 200
    payload = response.json()
    assert payload["active"]["provider"] == "api"
    assert payload["active"]["base_url"] == "https://api.example.test/v1"
    assert payload["active"]["model"] == "gpt-5.5"
    assert payload["active"]["has_api_key"] is True
    assert "secret-profile-key" not in response.text


def test_model_profile_endpoint_resolves_environment_secret_ref_without_exposing_key(monkeypatch):
    client = TestClient(app)
    monkeypatch.setenv("LLM_PROVIDER", "api")
    monkeypatch.setenv("OPENAI_BASE_URL", "https://api.example.test/v1")
    monkeypatch.setenv("MODEL_NAME", "gpt-5.5")
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.setenv("OPENAI_API_KEY_SECRET_REF", "prod-openai")
    monkeypatch.setenv("MODEL_SECRET_PROD_OPENAI", "server-side-secret-key")

    response = client.get("/api/model-profiles")

    assert response.status_code == 200
    payload = response.json()
    assert payload["active"]["has_api_key"] is True
    assert payload["active"]["api_key_secret_ref"] == "prod-openai"
    assert "server-side-secret-key" not in response.text


def test_model_profile_connection_test_uses_runtime_profile_without_exposing_key(monkeypatch):
    async def fake_test_connection(self):
        assert self.api_key == "runtime-secret-key"
        assert self.base_url == "https://api.example.test/v1"
        return ProviderResult(data={"status": "ok"}, provider="api", status="ok", model="gpt-5.5", warnings=[])

    monkeypatch.setattr("apps.api.llm_provider.LLMProvider.test_connection", fake_test_connection)
    client = TestClient(app)

    response = client.post(
        "/api/model-profiles/test",
        json={
            "profile": {
                "provider": "api",
                "base_url": "https://api.example.test/v1",
                "model": "gpt-5.5",
                "api_key": "runtime-secret-key",
                "temperature": 0.2,
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ok"] is True
    assert payload["profile"]["has_api_key"] is True
    assert payload["profile"]["base_url"] == "https://api.example.test/v1"
    assert "runtime-secret-key" not in response.text


def test_model_profile_connection_test_uses_server_secret_ref_without_exposing_key(monkeypatch):
    async def fake_test_connection(self):
        assert self.api_key == "server-side-runtime-key"
        assert self.base_url == "https://api.example.test/v1"
        return ProviderResult(data={"status": "ok"}, provider="api", status="ok", model="gpt-5.5", warnings=[])

    monkeypatch.setattr("apps.api.llm_provider.LLMProvider.test_connection", fake_test_connection)
    monkeypatch.setenv("MODEL_SECRET_STUDIO_PROD", "server-side-runtime-key")
    client = TestClient(app)

    response = client.post(
        "/api/model-profiles/test",
        json={
            "profile": {
                "provider": "api",
                "base_url": "https://api.example.test/v1",
                "model": "gpt-5.5",
                "api_key_secret_ref": "studio-prod",
                "temperature": 0.2,
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ok"] is True
    assert payload["profile"]["has_api_key"] is True
    assert payload["profile"]["api_key_secret_ref"] == "studio-prod"
    assert "server-side-runtime-key" not in response.text


def test_model_profile_connection_test_uses_short_timeout(monkeypatch):
    captured = {}

    async def fake_chat_json(self, prompt: str, system_prompt: str, max_tokens: int, timeout_seconds: float = 90):
        captured["timeout"] = timeout_seconds
        return ProviderResult(data=None, provider="local", status="fallback", model="fallback", warnings=["offline"])

    monkeypatch.setattr("apps.api.llm_provider.LLMProvider._chat_json", fake_chat_json)
    client = TestClient(app)

    response = client.post(
        "/api/model-profiles/test",
        json={"profile": {"provider": "local", "base_url": "http://127.0.0.1:11434/v1", "model": "qwen2.5:14b"}},
    )

    assert response.status_code == 200
    assert captured["timeout"] <= 8
    assert response.json()["ok"] is False


def test_api_rewrite_scene_updates_selected_scene():
    client = TestClient(app)
    generated = client.post("/api/generate", json={"text": SAMPLE_NOVEL, "title": "雨夜来信", "use_llm": False}).json()
    scene_id = generated["screenplay"]["scenes"][0]["id"]
    response = client.post(
        "/api/rewrite-scene",
        json={
            "screenplay": generated["screenplay"],
            "scene_id": scene_id,
            "instruction": "增强冲突",
            "mode": "strengthen_conflict",
            "use_llm": False,
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["validation"]["valid"] is True
    assert payload["changed_scene"]["id"] == scene_id
    assert payload["diff_summary"]
    assert payload["evidence"]
    assert payload["provider_status"] == "fallback"


def test_scene_evidence_retrieval_uses_source_refs_and_context():
    chapters = detect_chapters(SAMPLE_NOVEL)
    screenplay = build_fallback_screenplay(chapters, "雨夜来信", AdaptationStyle())
    scene = screenplay.scenes[0]
    evidence = retrieve_scene_evidence(screenplay.chapters, scene, "增强冲突")
    direct_ids = set(scene.source_refs[0].paragraph_ids)
    assert evidence
    assert direct_ids.intersection({item.paragraph_id for item in evidence})
    assert any(item.reason == "场景 source_refs 直接引用" for item in evidence)


def test_sqlite_fts_evidence_index_search(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    chapters = detect_chapters(SAMPLE_NOVEL)
    indexed = storage.index_project_evidence("proj_test", chapters)
    evidence = storage.search_project_evidence("proj_test", "钟楼 父亲 笔迹", limit=5)
    assert indexed == sum(len(chapter.paragraphs) for chapter in chapters)
    assert evidence
    assert any("钟楼" in item.text or "父亲" in item.text for item in evidence)
    assert evidence[0].reason == "FTS 证据库命中"


def test_sqlite_vector_evidence_search_finds_semantic_neighbors(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    chapters = detect_chapters(
        "第一章 雨夜\n\n林岚把旧信藏进抽屉，她担心父亲留下的暗号被人发现。\n\n"
        "第二章 追问\n\n陌生人逼近钟楼，所有人都在寻找那封信。\n\n"
        "第三章 交锋\n\n她必须决定是否公开真相。"
    )
    indexed = storage.index_project_evidence("proj_vector", chapters)
    evidence = storage.search_project_evidence("proj_vector", "保守秘密与亲人线索", limit=3)

    assert indexed == sum(len(chapter.paragraphs) for chapter in chapters)
    assert evidence
    assert evidence[0].reason == "向量证据库命中"
    assert "旧信" in evidence[0].text or "父亲" in evidence[0].text


def test_api_rewrite_scene_uses_llm_plan_when_available(monkeypatch):
    async def fake_rewrite_scene_object(self, prompt: str):
        return ProviderResult(
            data={
                "summary": "主角被迫在保守秘密和立刻行动之间做出选择。",
                "dramatic_purpose": "把原文中的线索压力转化成可表演的当场抉择。",
                "conflict": "如果主角继续沉默，关键证据会失效；如果开口，关系会立刻破裂。",
                "beats": ["线索出现", "对方逼问", "主角做出选择"],
                "emotional_shift": "从克制试探转向正面逼迫",
                "production_notes": ["用近景捕捉犹豫", "保留原文物件作为镜头线索"],
                "elements": [
                    {
                        "id": "el_001_002",
                        "text": "我可以继续装作不知道，但那封信今晚就会消失。",
                        "emotion": "压低声音",
                    }
                ],
            },
            provider="api",
            status="ok",
            model="test-model",
            warnings=[],
        )

    monkeypatch.setattr("apps.api.llm_provider.LLMProvider.rewrite_scene_object", fake_rewrite_scene_object)
    client = TestClient(app)
    generated = client.post("/api/generate", json={"text": SAMPLE_NOVEL, "title": "雨夜来信", "use_llm": False}).json()
    scene_id = generated["screenplay"]["scenes"][0]["id"]
    response = client.post(
        "/api/rewrite-scene",
        json={
            "screenplay": generated["screenplay"],
            "scene_id": scene_id,
            "instruction": "用证据增强人物选择代价",
            "mode": "strengthen_conflict",
            "use_llm": True,
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["provider_status"] == "ok"
    assert payload["validation"]["valid"] is True
    assert payload["changed_scene"]["summary"].startswith("主角被迫")
    assert payload["changed_scene"]["source_refs"] == generated["screenplay"]["scenes"][0]["source_refs"]
    assert payload["evidence"]


def test_project_job_persistence_flow(tmp_path, monkeypatch):
    from apps.api import storage
    import apps.api.jobs as jobs_module

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(jobs_module, "asyncio", jobs_module.asyncio)

    client = TestClient(app)
    created = client.post("/api/projects", json={"title": "企业项目", "source_text": SAMPLE_NOVEL}).json()
    project_id = created["id"]

    job = client.post("/api/jobs/generate", json={"project_id": project_id, "use_llm": False}).json()
    assert job["status"] in {"queued", "succeeded"}

    detail = client.get(f"/api/jobs/{job['id']}").json()
    assert detail["status"] == "succeeded"
    assert detail["progress"] == 100
    assert detail["events"]

    project = client.get(f"/api/projects/{project_id}").json()
    assert project["versions"]
    assert project["current_version_id"] == detail["result_version_id"]
    assert project["versions"][0]["screenplay"]["project"]["id"] == project_id

    projects = client.get("/api/projects").json()["projects"]
    assert any(item["id"] == project_id and item["version_count"] == 1 for item in projects)

    stream = client.get(f"/api/jobs/{job['id']}/events")
    assert stream.status_code == 200
    assert "event: job" in stream.text
    assert job["id"] in stream.text

    search = client.post(
        f"/api/projects/{project_id}/evidence/search",
        json={"query": "钟楼 父亲 笔迹", "limit": 5},
    )
    assert search.status_code == 200
    assert search.json()["evidence"]


def test_generation_job_runtime_model_profile_is_sanitized(tmp_path, monkeypatch):
    from apps.api import storage

    seen_profile = {}

    async def fake_generate_screenplay_object(self, prompt: str):
        seen_profile["api_key"] = self.api_key
        seen_profile["model"] = self.model
        return ProviderResult(data=None, provider=self.provider, status="fallback", model="fallback", warnings=["forced fallback"])

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr("apps.api.llm_provider.LLMProvider.generate_screenplay_object", fake_generate_screenplay_object)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "模型配置项目", "source_text": SAMPLE_NOVEL}).json()
    response = client.post(
        "/api/jobs/generate",
        json={
            "project_id": project["id"],
            "use_llm": True,
            "model_profile": {
                "provider": "api",
                "base_url": "https://api.example.test/v1",
                "model": "gpt-5.5",
                "api_key": "job-runtime-secret",
                "temperature": 0.3,
            },
        },
    )

    assert response.status_code == 200
    job = client.get(f"/api/jobs/{response.json()['id']}").json()
    project_detail = client.get(f"/api/projects/{project['id']}").json()

    assert seen_profile == {"api_key": "job-runtime-secret", "model": "gpt-5.5"}
    assert "job-runtime-secret" not in response.text
    assert "job-runtime-secret" not in str(job["request_payload"])
    assert job["request_payload"]["model_profile"]["has_api_key"] is True
    assert "api_key" not in job["request_payload"]["model_profile"]
    assert project_detail["settings"]["model_profile"]["has_api_key"] is True
    assert "job-runtime-secret" not in str(project_detail["settings"])


def test_generation_job_model_profile_secret_ref_is_resolved_and_sanitized(tmp_path, monkeypatch):
    from apps.api import storage

    seen_profile = {}

    async def fake_generate_screenplay_object(self, prompt: str):
        seen_profile["api_key"] = self.api_key
        seen_profile["model"] = self.model
        return ProviderResult(data=None, provider=self.provider, status="fallback", model="fallback", warnings=["forced fallback"])

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("MODEL_SECRET_STUDIO_PROD", "job-server-secret")
    monkeypatch.setattr("apps.api.llm_provider.LLMProvider.generate_screenplay_object", fake_generate_screenplay_object)

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "模型密钥引用项目", "source_text": SAMPLE_NOVEL}).json()
    response = client.post(
        "/api/jobs/generate",
        json={
            "project_id": project["id"],
            "use_llm": True,
            "model_profile": {
                "provider": "api",
                "base_url": "https://api.example.test/v1",
                "model": "gpt-5.5",
                "api_key_secret_ref": "studio-prod",
                "temperature": 0.3,
            },
        },
    )

    assert response.status_code == 200
    job = client.get(f"/api/jobs/{response.json()['id']}").json()
    project_detail = client.get(f"/api/projects/{project['id']}").json()

    assert seen_profile == {"api_key": "job-server-secret", "model": "gpt-5.5"}
    assert "job-server-secret" not in response.text
    assert "job-server-secret" not in str(job["request_payload"])
    assert job["request_payload"]["model_profile"]["has_api_key"] is True
    assert job["request_payload"]["model_profile"]["api_key_secret_ref"] == "studio-prod"
    assert "api_key" not in job["request_payload"]["model_profile"]
    assert project_detail["settings"]["model_profile"]["has_api_key"] is True
    assert project_detail["settings"]["model_profile"]["api_key_secret_ref"] == "studio-prod"
    assert "job-server-secret" not in str(project_detail["settings"])


def test_external_job_queue_requires_worker_to_execute(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "external")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "外部队列项目", "source_text": SAMPLE_NOVEL}).json()
    queued = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})
    assert queued.status_code == 200
    job = queued.json()
    assert job["status"] == "queued"
    assert job["queue_mode"] == "external"
    assert job["request_payload"]["use_llm"] is False
    assert client.get(f"/api/projects/{project['id']}").json()["versions"] == []

    status = client.get("/api/job-queue/status")
    assert status.status_code == 200
    assert status.json()["queued"] >= 1
    assert status.json()["mode"] == "external"

    worked = client.post("/api/workers/run-once")
    assert worked.status_code == 200
    assert worked.json()["job"]["status"] == "succeeded"
    detail = client.get(f"/api/jobs/{job['id']}").json()
    assert detail["status"] == "succeeded"
    assert detail["attempts"] == 1
    assert client.get(f"/api/projects/{project['id']}").json()["versions"]


def test_queued_job_can_be_canceled_before_worker_runs(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "external")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "取消队列项目", "source_text": SAMPLE_NOVEL}).json()
    job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()

    canceled = client.post(f"/api/jobs/{job['id']}", json={"actor": "项目负责人", "action": "cancel"})
    assert canceled.status_code == 200
    assert canceled.json()["status"] == "canceled"

    worked = client.post("/api/workers/run-once")
    assert worked.status_code == 200
    assert worked.json()["job"] is None
    detail = client.get(f"/api/jobs/{job['id']}").json()
    assert detail["status"] == "canceled"
    assert client.get(f"/api/projects/{project['id']}").json()["versions"] == []


def test_project_comments_create_audit_event(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "审阅项目", "source_text": SAMPLE_NOVEL}).json()
    response = client.post(
        f"/api/projects/{project['id']}/comments",
        json={
            "scene_id": "sc_001",
            "author": "制片审阅",
            "body": "这一场需要增强人物选择代价。",
            "status": "open",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["comment"]["scene_id"] == "sc_001"
    assert payload["comment"]["status"] == "open"
    assert payload["audit_event"]["event_type"] == "comment.created"

    detail = client.get(f"/api/projects/{project['id']}").json()
    assert detail["comments"][0]["body"].startswith("这一场")
    assert detail["audit_events"][0]["event_type"] == "comment.created"


def test_project_comment_status_update_creates_audit_event(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "审阅流转项目", "source_text": SAMPLE_NOVEL}).json()
    created = client.post(
        f"/api/projects/{project['id']}/comments",
        json={"scene_id": "sc_001", "author": "主编", "body": "需要补强这一场的证据链。"},
    ).json()

    response = client.patch(
        f"/api/projects/{project['id']}/comments/{created['comment']['id']}",
        json={"author": "主编", "status": "resolved"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["comment"]["status"] == "resolved"
    assert payload["audit_event"]["event_type"] == "comment.resolved"
    assert payload["audit_event"]["metadata"]["previous_status"] == "open"

    detail = client.get(f"/api/projects/{project['id']}").json()
    assert detail["comments"][0]["status"] == "resolved"
    assert detail["audit_events"][0]["event_type"] == "comment.resolved"


def test_project_review_assignment_reply_and_audit_filter(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "审阅增强项目", "source_text": SAMPLE_NOVEL}).json()
    created = client.post(
        f"/api/projects/{project['id']}/comments",
        json={
            "scene_id": "sc_001",
            "author": "主编",
            "assignee": "编剧",
            "body": "请补一版更有证据链的对白。",
        },
    )
    assert created.status_code == 200
    comment = created.json()["comment"]
    assert comment["assignee"] == "编剧"

    reply = client.post(
        f"/api/projects/{project['id']}/comments/{comment['id']}/replies",
        json={"author": "编剧", "body": "已根据原文证据补充对白。"},
    )
    assert reply.status_code == 200
    reply_payload = reply.json()
    assert reply_payload["comment"]["replies"][0]["body"].startswith("已根据")
    assert reply_payload["audit_event"]["event_type"] == "comment.replied"

    audit = client.get(f"/api/projects/{project['id']}/audit-events", params={"event_type": "comment.replied"})
    assert audit.status_code == 200
    events = audit.json()["audit_events"]
    assert len(events) == 1
    assert events[0]["event_type"] == "comment.replied"
    assert events[0]["metadata"]["comment_id"] == comment["id"]


def test_project_members_permissions_and_notifications(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "成员权限项目", "source_text": SAMPLE_NOVEL}).json()

    detail = client.get(f"/api/projects/{project['id']}").json()
    roles = {member["name"]: member["role"] for member in detail["members"]}
    assert roles["项目负责人"] == "owner"
    assert roles["审阅者"] == "reviewer"

    member = client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": "项目负责人", "name": "只读观察", "role": "viewer"},
    )
    assert member.status_code == 200
    assert member.json()["member"]["role"] == "viewer"

    blocked_member = client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": "审阅者", "name": "外部制片", "role": "producer"},
    )
    assert blocked_member.status_code == 403

    blocked_comment = client.post(
        f"/api/projects/{project['id']}/comments",
        json={"author": "只读观察", "body": "只读成员不能留下审阅意见。"},
    )
    assert blocked_comment.status_code == 403

    assigned = client.post(
        f"/api/projects/{project['id']}/comments",
        json={"author": "审阅者", "assignee": "编剧", "body": "请按证据链重写第一场。"},
    )
    assert assigned.status_code == 200
    comment = assigned.json()["comment"]
    assert assigned.json()["notifications"][0]["recipient"] == "编剧"
    assert assigned.json()["notifications"][0]["event_type"] == "comment.assigned"

    assignee_notifications = client.get(
        f"/api/projects/{project['id']}/notifications",
        params={"recipient": "编剧", "unread_only": "true"},
    )
    assert assignee_notifications.status_code == 200
    assert assignee_notifications.json()["notifications"][0]["unread"] is True

    reply = client.post(
        f"/api/projects/{project['id']}/comments/{comment['id']}/replies",
        json={"author": "编剧", "body": "已收到，会按证据链改。"},
    )
    assert reply.status_code == 200
    assert any(item["recipient"] == "审阅者" for item in reply.json()["notifications"])

    reviewer_notifications = client.get(
        f"/api/projects/{project['id']}/notifications",
        params={"recipient": "审阅者"},
    )
    assert reviewer_notifications.status_code == 200
    reviewer_notification = reviewer_notifications.json()["notifications"][0]
    assert reviewer_notification["event_type"] == "comment.replied"

    read = client.patch(
        f"/api/projects/{project['id']}/notifications/{reviewer_notification['id']}",
        json={"actor": "审阅者", "unread": False},
    )
    assert read.status_code == 200
    assert read.json()["notification"]["unread"] is False


def test_project_member_session_overrides_request_actor_for_permissions(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "会话权限项目", "source_text": SAMPLE_NOVEL}).json()

    session_response = client.post(
        "/api/auth/sessions",
        json={"project_id": project["id"], "member_name": "审阅者"},
    )
    assert session_response.status_code == 200
    session = session_response.json()
    assert session["member"]["name"] == "审阅者"
    headers = {"Authorization": f"Bearer {session['token']}"}

    blocked = client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": "项目负责人", "name": "越权成员", "role": "writer"},
        headers=headers,
    )
    assert blocked.status_code == 403

    comment = client.post(
        f"/api/projects/{project['id']}/comments",
        json={"author": "项目负责人", "body": "这里应该记录真实会话身份。"},
        headers=headers,
    )
    assert comment.status_code == 200
    payload = comment.json()
    assert payload["comment"]["author"] == "审阅者"
    assert payload["audit_event"]["actor"] == "审阅者"


def test_project_member_session_rejects_cross_project_token(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    first = client.post("/api/projects", json={"title": "项目一", "source_text": SAMPLE_NOVEL}).json()
    second = client.post("/api/projects", json={"title": "项目二", "source_text": SAMPLE_NOVEL}).json()
    session = client.post(
        "/api/auth/sessions",
        json={"project_id": first["id"], "member_name": "项目负责人"},
    ).json()

    response = client.post(
        f"/api/projects/{second['id']}/members",
        json={"name": "跨项目成员", "role": "writer"},
        headers={"Authorization": f"Bearer {session['token']}"},
    )

    assert response.status_code == 403


def test_project_batch_rewrite_job_creates_version_and_audit_event(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "批量改写项目", "source_text": SAMPLE_NOVEL}).json()
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    generated_detail = client.get(f"/api/jobs/{generated_job['id']}").json()
    assert generated_detail["status"] == "succeeded"

    first_version = client.get(f"/api/projects/{project['id']}").json()["versions"][0]
    scene_id = first_version["screenplay"]["scenes"][0]["id"]

    rewrite_job = client.post(
        "/api/jobs/rewrite",
        json={
            "project_id": project["id"],
            "actor": "编剧",
            "scene_ids": [scene_id],
            "mode": "short_drama",
            "instruction": "批量强化短剧钩子，但保留原文证据链。",
            "use_llm": False,
            "max_scenes": 3,
        },
    )

    assert rewrite_job.status_code == 200
    job_payload = rewrite_job.json()
    assert job_payload["kind"] == "rewrite"

    detail = client.get(f"/api/jobs/{job_payload['id']}").json()
    assert detail["status"] == "succeeded"
    assert detail["progress"] == 100
    assert detail["result_version_id"]
    assert any(event["stage_id"] == "rewrite.completed" for event in detail["events"])

    rewritten_project = client.get(f"/api/projects/{project['id']}").json()
    assert len(rewritten_project["versions"]) == 2
    latest = rewritten_project["versions"][-1]
    assert latest["id"] == detail["result_version_id"]
    assert latest["label"].startswith("批量改写版本")
    changed_scene = next(scene for scene in latest["screenplay"]["scenes"] if scene["id"] == scene_id)
    assert "强钩子" in changed_scene["summary"]
    assert changed_scene["source_refs"] == first_version["screenplay"]["scenes"][0]["source_refs"]
    assert rewritten_project["audit_events"][0]["event_type"] == "rewrite.batch_completed"
    assert rewritten_project["audit_events"][0]["metadata"]["scene_ids"] == scene_id


def test_project_version_history_compare_and_restore(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "版本治理项目", "source_text": SAMPLE_NOVEL}).json()
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    first_detail = client.get(f"/api/jobs/{generated_job['id']}").json()
    first_version_id = first_detail["result_version_id"]
    first_version = client.get(f"/api/projects/{project['id']}").json()["versions"][0]
    scene_id = first_version["screenplay"]["scenes"][0]["id"]

    rewrite_job = client.post(
        "/api/jobs/rewrite",
        json={
            "project_id": project["id"],
            "actor": "编剧",
            "scene_ids": [scene_id],
            "mode": "short_drama",
            "instruction": "把第一个场景推向更强钩子。",
            "use_llm": False,
            "max_scenes": 1,
        },
    ).json()
    second_detail = client.get(f"/api/jobs/{rewrite_job['id']}").json()
    second_version_id = second_detail["result_version_id"]

    versions = client.get(f"/api/projects/{project['id']}/versions")
    assert versions.status_code == 200
    version_payload = versions.json()["versions"]
    assert [item["id"] for item in version_payload] == [second_version_id, first_version_id]
    assert version_payload[0]["is_current"] is True
    assert version_payload[0]["scene_count"] == 3
    assert version_payload[0]["yaml_bytes"] > 0

    compare = client.get(
        f"/api/projects/{project['id']}/versions/compare",
        params={"base_version_id": first_version_id, "target_version_id": second_version_id},
    )
    assert compare.status_code == 200
    compare_payload = compare.json()
    assert compare_payload["base_version_id"] == first_version_id
    assert compare_payload["target_version_id"] == second_version_id
    assert compare_payload["changed_scenes"]
    assert any(item["scene_id"] == scene_id for item in compare_payload["changed_scenes"])
    assert compare_payload["yaml_diff_preview"]

    restored = client.post(
        f"/api/projects/{project['id']}/versions/{first_version_id}/restore",
        json={"actor": "项目负责人"},
    )
    assert restored.status_code == 200
    restored_project = restored.json()
    assert restored_project["current_version_id"] == first_version_id
    assert restored_project["audit_events"][0]["event_type"] == "version.restored"
    assert restored_project["audit_events"][0]["metadata"]["version_id"] == first_version_id


def test_project_version_restore_requires_admin_permission(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "版本权限项目", "source_text": SAMPLE_NOVEL}).json()
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    version_id = client.get(f"/api/jobs/{generated_job['id']}").json()["result_version_id"]

    blocked = client.post(
        f"/api/projects/{project['id']}/versions/{version_id}/restore",
        json={"actor": "审阅者"},
    )

    assert blocked.status_code == 403


def test_project_readiness_blocks_empty_project(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "空项目", "source_text": ""}).json()

    response = client.get(f"/api/projects/{project['id']}/readiness")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "blocked"
    assert payload["score"] < 70
    blocker_ids = {item["id"] for item in payload["blockers"]}
    assert {"source", "version"}.issubset(blocker_ids)
    assert payload["next_actions"]


def test_project_readiness_tracks_quality_review_and_import_security(tmp_path, monkeypatch):
    from apps.api import main as main_module
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setattr(main_module, "IMPORT_UPLOAD_DIR", data_dir / "imports")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "交付门禁项目", "source_text": ""}).json()
    import_job = client.post(
        "/api/jobs/import-source",
        data={"project_id": project["id"], "actor": "编剧"},
        files={"file": ("rain-letter.md", SAMPLE_NOVEL.encode("utf-8"), "text/markdown")},
    ).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})
    client.post(
        f"/api/projects/{project['id']}/comments",
        json={"author": "主编", "body": "上线前需要复核第一场证据链。"},
    )

    response = client.get(f"/api/projects/{project['id']}/readiness")

    assert response.status_code == 200
    payload = response.json()
    assert payload["current_version_id"]
    assert payload["status"] == "warning"
    assert payload["score"] >= 70
    passed_ids = {item["id"] for item in payload["passed"]}
    warning_ids = {item["id"] for item in payload["warnings"]}
    assert {"source", "version", "schema", "quality", "import_security", "jobs"}.issubset(passed_ids)
    assert "review" in warning_ids
    import_security = next(item for item in payload["passed"] if item["id"] == "import_security")
    assert import_security["evidence"]["security_verdict"] == "clean"
    assert import_security["evidence"]["sha256"] == import_job["result_payload"]["sha256"]


def test_project_export_governance_blocks_empty_project_and_audits(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Export blocked", "source_text": ""}).json()

    response = client.post(
        f"/api/projects/{project['id']}/exports",
        json={"format": "fountain", "actor": "制片", "enforce_readiness": True},
    )

    assert response.status_code == 409
    detail = response.json()["detail"]
    assert detail["export"]["status"] == "blocked"
    assert detail["export"]["format"] == "fountain"
    assert detail["readiness"]["status"] == "blocked"
    assert detail["audit_event"]["event_type"] == "export.blocked"

    history = client.get(f"/api/projects/{project['id']}/exports").json()
    assert history["blocked"] == 1
    assert history["exports"][0]["audit_event_id"] == detail["audit_event"]["id"]
    refreshed = client.get(f"/api/projects/{project['id']}").json()
    assert refreshed["audit_events"][0]["event_type"] == "export.blocked"


def test_project_export_governance_records_version_readiness_and_hash(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Export ready", "source_text": SAMPLE_NOVEL}).json()
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    version_id = client.get(f"/api/jobs/{generated_job['id']}").json()["result_version_id"]

    response = client.post(
        f"/api/projects/{project['id']}/exports",
        json={"format": "markdown", "actor": "制片", "enforce_readiness": True},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["export"]["status"] == "succeeded"
    assert payload["export"]["version_id"] == version_id
    assert payload["export"]["format"] == "markdown"
    assert payload["export"]["size_bytes"] == len(payload["content"].encode("utf-8"))
    assert len(payload["export"]["sha256"]) == 64
    assert payload["export"]["readiness_score"] == payload["readiness"]["score"]
    assert payload["audit_event"]["event_type"] == "export.created"
    assert "# Export ready" in payload["content"]

    history = client.get(f"/api/projects/{project['id']}/exports").json()
    assert history["succeeded"] == 1
    assert history["exports"][0]["sha256"] == payload["export"]["sha256"]
    refreshed = client.get(f"/api/projects/{project['id']}").json()
    assert refreshed["exports"][0]["id"] == payload["export"]["id"]
    assert refreshed["audit_events"][0]["metadata"]["export_id"] == payload["export"]["id"]


def test_project_export_requires_export_permission(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Export permission", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})

    response = client.post(
        f"/api/projects/{project['id']}/exports",
        json={"format": "yaml", "actor": "审阅者"},
    )

    assert response.status_code == 403


def test_project_approval_writer_submit_and_history(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Approval submit", "source_text": SAMPLE_NOVEL}).json()
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    version_id = client.get(f"/api/jobs/{generated_job['id']}").json()["result_version_id"]

    response = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={
            "actor": "编剧",
            "note": "请按当前版本进入交付审批。",
            "requested_export_format": "fountain",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["approval"]["status"] == "submitted"
    assert payload["approval"]["version_id"] == version_id
    assert payload["approval"]["submitted_by"] == "编剧"
    assert payload["approval"]["requested_export_format"] == "fountain"
    assert payload["approval"]["readiness_status"] in {"ready", "warning"}
    assert payload["audit_event"]["event_type"] == "approval.submitted"
    assert any(item["recipient"] == "制片" for item in payload["notifications"])

    history = client.get(f"/api/projects/{project['id']}/approvals").json()
    assert history["pending"] == 1
    assert history["approvals"][0]["id"] == payload["approval"]["id"]
    refreshed = client.get(f"/api/projects/{project['id']}").json()
    assert refreshed["approvals"][0]["audit_event_id"] == payload["audit_event"]["id"]
    assert refreshed["audit_events"][0]["event_type"] == "approval.submitted"


def test_project_approval_producer_can_approve_and_notify_submitter(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Approval approve", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})
    submitted = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "编剧", "note": "当前版本可进入制片交付。"},
    ).json()

    response = client.post(
        f"/api/projects/{project['id']}/approvals/{submitted['approval']['id']}/decision",
        json={"actor": "制片", "decision": "approve", "note": "同意作为交付版本。"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["approval"]["status"] == "approved"
    assert payload["approval"]["decided_by"] == "制片"
    assert payload["approval"]["decision_note"] == "同意作为交付版本。"
    assert payload["audit_event"]["event_type"] == "approval.approved"
    assert payload["notifications"][0]["recipient"] == "编剧"

    history = client.get(f"/api/projects/{project['id']}/approvals").json()
    assert history["approved"] == 1
    assert history["pending"] == 0
    refreshed = client.get(f"/api/projects/{project['id']}").json()
    assert refreshed["approvals"][0]["decision_audit_event_id"] == payload["audit_event"]["id"]


def test_project_approval_reviewer_cannot_submit_or_decide(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Approval permission", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})

    blocked_submit = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "审阅者", "note": "审阅者不能提交交付审批。"},
    )
    assert blocked_submit.status_code == 403

    submitted = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "编剧", "note": "提交给制片复核。"},
    ).json()
    blocked_decision = client.post(
        f"/api/projects/{project['id']}/approvals/{submitted['approval']['id']}/decision",
        json={"actor": "审阅者", "decision": "approve", "note": "审阅者不能批准。"},
    )
    assert blocked_decision.status_code == 403


def test_project_approval_blocks_empty_project_and_audits(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Approval blocked", "source_text": ""}).json()

    response = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "编剧", "note": "空项目不能提交交付审批。"},
    )

    assert response.status_code == 409
    detail = response.json()["detail"]
    assert detail["approval"]["status"] == "blocked"
    assert detail["approval"]["version_id"] is None
    assert detail["readiness"]["status"] == "blocked"
    assert detail["audit_event"]["event_type"] == "approval.blocked"

    history = client.get(f"/api/projects/{project['id']}/approvals").json()
    assert history["blocked"] == 1
    assert history["approvals"][0]["audit_event_id"] == detail["audit_event"]["id"]


def test_project_delivery_package_requires_approved_approval_and_records_manifest(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Delivery package ready", "source_text": SAMPLE_NOVEL}).json()
    generated_job = client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False}).json()
    version_id = client.get(f"/api/jobs/{generated_job['id']}").json()["result_version_id"]
    submitted = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "编剧", "note": "请求进入正式交付包。"},
    ).json()
    approved = client.post(
        f"/api/projects/{project['id']}/approvals/{submitted['approval']['id']}/decision",
        json={"actor": "制片", "decision": "approve", "note": "批准打包。"},
    ).json()

    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "制片",
            "approval_id": approved["approval"]["id"],
            "formats": ["yaml", "markdown", "fountain"],
            "note": "正式交付给制片组。",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    package = payload["package"]
    assert package["status"] == "succeeded"
    assert package["version_id"] == version_id
    assert package["approval_id"] == approved["approval"]["id"]
    assert package["formats"] == ["yaml", "markdown", "fountain"]
    assert len(package["assets"]) == 4
    assert len(payload["assets"]) == 4
    assert package["total_size_bytes"] == sum(asset["size_bytes"] for asset in payload["assets"])
    assert package["artifact_count"] == 4
    assert len(package["manifest_sha256"]) == 64
    manifest = next(asset for asset in payload["assets"] if asset["filename"].endswith("delivery-manifest.json"))
    assert manifest["sha256"] == package["manifest_sha256"]
    manifest_payload = __import__("json").loads(manifest["content"])
    assert manifest_payload["schema"] == "scriptbridge.delivery-package.v1"
    assert manifest_payload["package_id"] == package["id"]
    assert manifest_payload["approval"]["id"] == approved["approval"]["id"]
    assert len(manifest_payload["assets"]) == 3
    assert payload["audit_event"]["event_type"] == "delivery_package.created"

    history = client.get(f"/api/projects/{project['id']}/delivery-packages").json()
    assert history["succeeded"] == 1
    assert history["packages"][0]["manifest_sha256"] == package["manifest_sha256"]
    refreshed = client.get(f"/api/projects/{project['id']}").json()
    assert refreshed["delivery_packages"][0]["id"] == package["id"]
    assert refreshed["audit_events"][0]["metadata"]["package_id"] == package["id"]


def test_project_delivery_package_assets_are_signed_and_downloadable(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_SIGNING_SECRET", "test-delivery-secret")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Downloadable package", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})
    submitted = client.post(
        f"/api/projects/{project['id']}/approvals",
        json={"actor": "编剧", "note": "请求下载中心验证。"},
    ).json()
    approved = client.post(
        f"/api/projects/{project['id']}/approvals/{submitted['approval']['id']}/decision",
        json={"actor": "制片", "decision": "approve", "note": "批准下载中心验证。"},
    ).json()

    created = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "制片",
            "approval_id": approved["approval"]["id"],
            "formats": ["yaml", "markdown"],
            "note": "生成可复取交付包。",
        },
    ).json()

    package = created["package"]
    assert package["storage_provider"] == "local"
    assert package["artifact_count"] == 3
    assert package["download_expires_at"]
    assert len(package["assets"]) == 3
    assert all(asset["storage_key"] for asset in package["assets"])
    assert all(asset["download_url"] for asset in package["assets"])
    assert all(asset["content"] is None for asset in package["assets"])
    assert len(created["assets"]) == 3
    assert all(asset["content"] for asset in created["assets"])

    yaml_asset = next(asset for asset in package["assets"] if asset["filename"].endswith(".yaml"))
    downloaded = client.get(yaml_asset["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["content-type"].startswith("application/x-yaml")
    assert downloaded.content
    assert __import__("hashlib").sha256(downloaded.content).hexdigest() == yaml_asset["sha256"]
    assert "filename*" in downloaded.headers["content-disposition"]

    history = client.get(f"/api/projects/{project['id']}/delivery-packages?actor=制片").json()
    history_asset = history["packages"][0]["assets"][0]
    assert history_asset["download_url"]
    assert history["packages"][0]["artifact_count"] == 3

    reviewer_history = client.get(f"/api/projects/{project['id']}/delivery-packages?actor=审阅者")
    assert reviewer_history.status_code == 403

    tampered_url = yaml_asset["download_url"].replace("token=", "token=x", 1)
    tampered = client.get(tampered_url)
    assert tampered.status_code == 403


def _disabled_bad_encoded_delivery_package_provider_block():
    pass


def _disabled_bad_encoded_delivery_package_provider_block_legacy():
    """
    data_dir = tmp_path / "scriptbridge_data"
    artifact_root = tmp_path / "artifact-root"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "local")
    monkeypatch.setenv("DELIVERY_ARTIFACT_LOCAL_DIR", str(artifact_root))
    monkeypatch.setenv("DELIVERY_SIGNING_SECRET", "test-delivery-secret")

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "Configured artifact root")
    created = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "鍒剁墖",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "楠岃瘉 artifact provider 鏍圭洰褰曘€?,
        },
    )

    assert created.status_code == 200
    package = created.json()["package"]
    assert package["storage_provider"] == "local"
    assert package["artifact_count"] == 2
    yaml_asset = next(asset for asset in package["assets"] if asset["filename"].endswith(".yaml"))
    assert yaml_asset["storage_provider"] == "local"
    assert (artifact_root / yaml_asset["storage_key"]).exists()
    assert not (data_dir / "delivery-artifacts" / yaml_asset["storage_key"]).exists()

    downloaded = client.get(yaml_asset["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["x-content-sha256"] == yaml_asset["sha256"]


def test_project_delivery_package_remote_provider_missing_config_returns_503(tmp_path, monkeypatch):
    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "s3")

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "Remote provider package")
    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "鍒剁墖",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "杩滅瀵硅薄瀛樺偍鏈帴 SDK 鏃跺簲鏄庣‘澶辫触銆?,
        },
    )

    assert response.status_code == 503
    assert "s3" in response.json()["detail"]
    history = client.get(f"/api/projects/{project['id']}/delivery-packages?actor=鍒剁墖").json()
    assert history["total"] == 0


def test_project_delivery_package_blocks_without_approved_approval(tmp_path, monkeypatch):
    """
    pass


def test_project_delivery_package_uses_configured_local_artifact_provider_root(tmp_path, monkeypatch):
    data_dir = tmp_path / "scriptbridge_data"
    artifact_root = tmp_path / "artifact-root"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "local")
    monkeypatch.setenv("DELIVERY_ARTIFACT_LOCAL_DIR", str(artifact_root))
    monkeypatch.setenv("DELIVERY_SIGNING_SECRET", "test-delivery-secret")

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "Configured artifact root")
    created = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "producer_ascii",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "verify artifact provider root",
        },
    )

    assert created.status_code == 200
    package = created.json()["package"]
    assert package["storage_provider"] == "local"
    assert package["artifact_count"] == 2
    yaml_asset = next(asset for asset in package["assets"] if asset["filename"].endswith(".yaml"))
    assert yaml_asset["storage_provider"] == "local"
    assert (artifact_root / yaml_asset["storage_key"]).exists()
    assert not (data_dir / "delivery-artifacts" / yaml_asset["storage_key"]).exists()

    downloaded = client.get(yaml_asset["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["x-content-sha256"] == yaml_asset["sha256"]


def test_local_artifact_provider_rejects_prefix_sibling_path(tmp_path):
    from apps.api.artifact_storage import LocalArtifactStorageProvider

    provider = LocalArtifactStorageProvider(root=tmp_path / "artifact-root")
    sibling = tmp_path / "artifact-root2" / "escaped.yaml"

    try:
        provider.write("../artifact-root2/escaped.yaml", b"escaped")
    except ValueError:
        pass
    else:
        raise AssertionError("Expected local artifact provider to reject sibling path escape.")

    assert not sibling.exists()


class _FakeS3Body:
    def __init__(self, content: bytes):
        self._content = content

    def read(self) -> bytes:
        return self._content


class _FakeS3Client:
    def __init__(self):
        self.objects = {}
        self.put_calls = []
        self.get_calls = []

    def put_object(self, **kwargs):
        self.put_calls.append(kwargs)
        self.objects[(kwargs["Bucket"], kwargs["Key"])] = kwargs["Body"]

    def get_object(self, **kwargs):
        self.get_calls.append(kwargs)
        key = (kwargs["Bucket"], kwargs["Key"])
        if key not in self.objects:
            raise KeyError(kwargs["Key"])
        return {"Body": _FakeS3Body(self.objects[key])}


class _FakeQiniuInfo:
    def __init__(self, status_code: int = 200, text_body: str = ""):
        self.status_code = status_code
        self.text_body = text_body


class _FakeQiniuClient:
    def __init__(self):
        self.objects = {}
        self.upload_tokens = []
        self.put_calls = []
        self.download_urls = []

    def upload_token(self, bucket: str, key: str, expires: int) -> str:
        token = f"upload-token:{bucket}:{key}:{expires}"
        self.upload_tokens.append({"bucket": bucket, "key": key, "expires": expires, "token": token})
        return token

    def put_data(self, upload_token: str, key: str, content: bytes):
        self.put_calls.append({"upload_token": upload_token, "key": key, "content": content})
        self.objects[key] = content
        return {"key": key}, _FakeQiniuInfo()

    def private_download_url(self, base_url: str, expires: int) -> str:
        signed_url = f"{base_url}?signed={expires}"
        self.download_urls.append({"base_url": base_url, "expires": expires, "signed_url": signed_url})
        return signed_url


class _FakeQiniuHttpResponse:
    def __init__(self, content: bytes = b"", status_code: int = 200, text: str = ""):
        self.content = content
        self.status_code = status_code
        self.text = text


def test_s3_artifact_provider_writes_and_reads_with_prefix():
    from apps.api.artifact_storage import S3ArtifactStorageProvider

    client = _FakeS3Client()
    provider = S3ArtifactStorageProvider(
        name="s3",
        client=client,
        bucket="delivery-bucket",
        key_prefix="scriptbridge/prod",
    )

    provider.write("proj_1/package_1/screenplay.yaml", b"screenplay-yaml")
    content = provider.read("proj_1/package_1/screenplay.yaml")

    assert content == b"screenplay-yaml"
    assert client.put_calls[0]["Bucket"] == "delivery-bucket"
    assert client.put_calls[0]["Key"] == "scriptbridge/prod/proj_1/package_1/screenplay.yaml"
    assert client.get_calls[0]["Key"] == "scriptbridge/prod/proj_1/package_1/screenplay.yaml"


def test_qiniu_artifact_provider_writes_and_reads_with_prefix():
    from apps.api.artifact_storage import QiniuArtifactStorageProvider

    sdk_client = _FakeQiniuClient()

    def fake_http_get(url: str):
        assert url == "https://assets.example.test/scriptbridge/prod/proj_1/package_1/screenplay.yaml?signed=900"
        return _FakeQiniuHttpResponse(sdk_client.objects["scriptbridge/prod/proj_1/package_1/screenplay.yaml"])

    provider = QiniuArtifactStorageProvider(
        name="qiniu",
        client=sdk_client,
        bucket="delivery-bucket",
        domain="assets.example.test",
        key_prefix="scriptbridge/prod",
        upload_token_expires=600,
        download_url_expires=900,
        http_get=fake_http_get,
    )

    provider.write("proj_1/package_1/screenplay.yaml", b"screenplay-yaml")
    content = provider.read("proj_1/package_1/screenplay.yaml")

    assert content == b"screenplay-yaml"
    assert sdk_client.upload_tokens[0]["bucket"] == "delivery-bucket"
    assert sdk_client.upload_tokens[0]["key"] == "scriptbridge/prod/proj_1/package_1/screenplay.yaml"
    assert sdk_client.upload_tokens[0]["expires"] == 600
    assert sdk_client.put_calls[0]["key"] == "scriptbridge/prod/proj_1/package_1/screenplay.yaml"
    assert sdk_client.download_urls[0]["base_url"] == (
        "https://assets.example.test/scriptbridge/prod/proj_1/package_1/screenplay.yaml"
    )


def test_s3_artifact_provider_rejects_drive_and_colon_keys():
    from apps.api.artifact_storage import S3ArtifactStorageProvider

    provider = S3ArtifactStorageProvider(
        name="s3",
        client=_FakeS3Client(),
        bucket="delivery-bucket",
    )

    for storage_key in ["C:\\temp\\screenplay.yaml", "proj_1/package:1/screenplay.yaml"]:
        try:
            provider.write(storage_key, b"screenplay-yaml")
        except ValueError:
            pass
        else:
            raise AssertionError(f"Expected S3 artifact provider to reject unsafe key: {storage_key}")


def test_delivery_package_can_use_configured_s3_provider_with_fake_client(tmp_path, monkeypatch):
    from apps.api import artifact_storage

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    fake_client = _FakeS3Client()
    captured_config = {}

    def fake_client_factory(**kwargs):
        captured_config.update(kwargs)
        return fake_client

    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "s3")
    monkeypatch.setenv("DELIVERY_ARTIFACT_ENDPOINT", "https://s3.example.test")
    monkeypatch.setenv("DELIVERY_ARTIFACT_BUCKET", "delivery-bucket")
    monkeypatch.setenv("DELIVERY_ARTIFACT_ACCESS_KEY_ID", "test-access-key")
    monkeypatch.setenv("DELIVERY_ARTIFACT_SECRET_ACCESS_KEY", "test-secret-key")
    monkeypatch.setenv("DELIVERY_ARTIFACT_REGION", "ap-northeast-1")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PREFIX", "scriptbridge/prod/")
    monkeypatch.setattr(artifact_storage, "S3_CLIENT_FACTORY", fake_client_factory)

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "S3 artifact package")
    created = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "producer_ascii",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "write delivery package to S3-compatible storage",
        },
    )

    assert created.status_code == 200
    package = created.json()["package"]
    assert package["storage_provider"] == "s3"
    assert package["artifact_count"] == 2
    assert captured_config["endpoint_url"] == "https://s3.example.test"
    assert captured_config["region_name"] == "ap-northeast-1"
    assert captured_config["aws_access_key_id"] == "test-access-key"
    assert captured_config["aws_secret_access_key"] == "test-secret-key"
    assert all(asset["storage_provider"] == "s3" for asset in package["assets"])
    assert all(key[0] == "delivery-bucket" for key in fake_client.objects)
    assert all(key[1].startswith("scriptbridge/prod/") for key in fake_client.objects)

    yaml_asset = next(asset for asset in package["assets"] if asset["filename"].endswith(".yaml"))
    downloaded = client.get(yaml_asset["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["x-content-sha256"] == yaml_asset["sha256"]
    assert __import__("hashlib").sha256(downloaded.content).hexdigest() == yaml_asset["sha256"]


def test_delivery_package_s3_provider_requires_credentials(tmp_path, monkeypatch):
    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "s3")
    monkeypatch.setenv("DELIVERY_ARTIFACT_BUCKET", "delivery-bucket")

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "S3 missing credentials")
    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "producer_ascii",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "missing s3 credentials should fail",
        },
    )

    assert response.status_code == 503
    detail = response.json()["detail"]
    assert "DELIVERY_ARTIFACT_ACCESS_KEY_ID" in detail
    assert "DELIVERY_ARTIFACT_SECRET_ACCESS_KEY" in detail
    history = client.get(f"/api/projects/{project['id']}/delivery-packages?actor=producer_ascii").json()
    assert history["total"] == 0


def test_minio_artifact_provider_uses_path_style_config(monkeypatch):
    from apps.api import artifact_storage

    captured_config = {}

    def fake_client_factory(**kwargs):
        captured_config.update(kwargs)
        return _FakeS3Client()

    monkeypatch.setenv("DELIVERY_ARTIFACT_ENDPOINT", "http://127.0.0.1:9000")
    monkeypatch.setenv("DELIVERY_ARTIFACT_BUCKET", "delivery-bucket")
    monkeypatch.setenv("DELIVERY_ARTIFACT_ACCESS_KEY_ID", "minio-access")
    monkeypatch.setenv("DELIVERY_ARTIFACT_SECRET_ACCESS_KEY", "minio-secret")
    monkeypatch.setattr(artifact_storage, "S3_CLIENT_FACTORY", fake_client_factory)

    provider = artifact_storage.delivery_artifact_provider(__import__("pathlib").Path("."), "minio")

    assert provider.name == "minio"
    assert captured_config["endpoint_url"] == "http://127.0.0.1:9000"
    assert captured_config["config"].s3["addressing_style"] == "path"


def test_delivery_package_can_use_configured_qiniu_provider_with_fake_client(tmp_path, monkeypatch):
    from apps.api import artifact_storage

    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    fake_client = _FakeQiniuClient()

    def fake_client_factory(access_key: str, secret_key: str):
        assert access_key == "qiniu-access"
        assert secret_key == "qiniu-secret"
        return fake_client

    def fake_http_get(url: str):
        path = url.split("assets.example.test/", 1)[1].split("?signed=", 1)[0]
        return _FakeQiniuHttpResponse(fake_client.objects[path])

    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "qiniu")
    monkeypatch.setenv("DELIVERY_ARTIFACT_BUCKET", "delivery-bucket")
    monkeypatch.setenv("DELIVERY_ARTIFACT_QINIU_ACCESS_KEY", "qiniu-access")
    monkeypatch.setenv("DELIVERY_ARTIFACT_QINIU_SECRET_KEY", "qiniu-secret")
    monkeypatch.setenv("DELIVERY_ARTIFACT_QINIU_DOMAIN", "assets.example.test")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PREFIX", "scriptbridge/prod")
    monkeypatch.setattr(artifact_storage, "QINIU_CLIENT_FACTORY", fake_client_factory)
    monkeypatch.setattr(artifact_storage, "QINIU_HTTP_GET", fake_http_get)

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "Qiniu artifact package")
    created = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "producer_ascii",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "write delivery package to Qiniu Kodo storage",
        },
    )

    assert created.status_code == 200
    package = created.json()["package"]
    assert package["storage_provider"] == "qiniu"
    assert package["artifact_count"] == 2
    assert all(asset["storage_provider"] == "qiniu" for asset in package["assets"])
    assert all(key.startswith("scriptbridge/prod/") for key in fake_client.objects)

    yaml_asset = next(asset for asset in package["assets"] if asset["filename"].endswith(".yaml"))
    downloaded = client.get(yaml_asset["download_url"])
    assert downloaded.status_code == 200
    assert downloaded.headers["x-content-sha256"] == yaml_asset["sha256"]
    assert __import__("hashlib").sha256(downloaded.content).hexdigest() == yaml_asset["sha256"]


def test_delivery_package_qiniu_provider_requires_configuration(tmp_path, monkeypatch):
    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "qiniu")
    monkeypatch.setenv("DELIVERY_ARTIFACT_BUCKET", "delivery-bucket")

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "Qiniu missing config")
    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "producer_ascii",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "missing qiniu credentials should fail",
        },
    )

    assert response.status_code == 503
    detail = response.json()["detail"]
    assert "qiniu" in detail
    assert "DELIVERY_ARTIFACT_QINIU_ACCESS_KEY" in detail
    assert "DELIVERY_ARTIFACT_QINIU_SECRET_KEY" in detail
    assert "DELIVERY_ARTIFACT_QINIU_DOMAIN" in detail
    history = client.get(f"/api/projects/{project['id']}/delivery-packages?actor=producer_ascii").json()
    assert history["total"] == 0


def test_project_delivery_package_remote_provider_missing_config_returns_503_legacy(tmp_path, monkeypatch):
    data_dir = tmp_path / "scriptbridge_data"
    _configure_test_storage(monkeypatch, data_dir)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")
    monkeypatch.setenv("DELIVERY_ARTIFACT_PROVIDER", "s3")

    client = TestClient(app)
    project, _, approval = _approved_delivery_project(client, "Remote provider package")
    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={
            "actor": "producer_ascii",
            "approval_id": approval["id"],
            "formats": ["yaml"],
            "note": "remote provider should fail without required config",
        },
    )

    assert response.status_code == 503
    assert "s3" in response.json()["detail"]
    history = client.get(f"/api/projects/{project['id']}/delivery-packages?actor=producer_ascii").json()
    assert history["total"] == 0


def test_project_delivery_package_blocks_without_approved_approval(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Delivery package blocked", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})

    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={"actor": "制片", "formats": ["yaml"], "require_approval": True},
    )

    assert response.status_code == 409
    detail = response.json()["detail"]
    assert detail["package"]["status"] == "blocked"
    assert detail["package"]["assets"] == []
    assert "已批准" in detail["package"]["blockers"][0]
    assert detail["audit_event"]["event_type"] == "delivery_package.blocked"

    history = client.get(f"/api/projects/{project['id']}/delivery-packages").json()
    assert history["blocked"] == 1
    assert history["packages"][0]["audit_event_id"] == detail["audit_event"]["id"]


def test_project_delivery_package_requires_producer_permission(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "Delivery package permission", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})

    response = client.post(
        f"/api/projects/{project['id']}/delivery-packages",
        json={"actor": "审阅者", "formats": ["yaml"], "require_approval": False},
    )

    assert response.status_code == 403


def test_project_batch_rewrite_job_requires_rewrite_permission(tmp_path, monkeypatch):
    from apps.api import storage

    _enable_strict_rbac(monkeypatch)
    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "只读批量改写项目", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})
    member = client.post(
        f"/api/projects/{project['id']}/members",
        json={"actor": "项目负责人", "name": "只读观察", "role": "viewer"},
    )
    assert member.status_code == 200

    blocked = client.post(
        "/api/jobs/rewrite",
        json={
            "project_id": project["id"],
            "actor": "只读观察",
            "mode": "tighten_evidence",
            "instruction": "尝试批量改写。",
            "use_llm": False,
        },
    )

    assert blocked.status_code == 403


def test_local_single_user_mode_allows_reviewer_actor_to_start_batch_rewrite(tmp_path, monkeypatch):
    from apps.api import storage

    data_dir = tmp_path / "scriptbridge_data"
    monkeypatch.setattr(storage, "DATA_DIR", data_dir)
    monkeypatch.setattr(storage, "PROJECTS_DIR", data_dir / "projects")
    monkeypatch.setattr(storage, "JOBS_DIR", data_dir / "jobs")
    monkeypatch.setattr(storage, "DB_PATH", data_dir / "scriptbridge.sqlite3")
    monkeypatch.delenv("AUTH_MODE", raising=False)
    monkeypatch.setenv("JOB_QUEUE_MODE", "inline")

    client = TestClient(app)
    project = client.post("/api/projects", json={"title": "本地全权限改写", "source_text": SAMPLE_NOVEL}).json()
    client.post("/api/jobs/generate", json={"project_id": project["id"], "use_llm": False})

    response = client.post(
        "/api/jobs/rewrite",
        json={
            "project_id": project["id"],
            "actor": "审阅者",
            "mode": "tighten_evidence",
            "instruction": "本地单用户模式下直接启动批量改写。",
            "use_llm": False,
            "max_scenes": 1,
        },
    )

    assert response.status_code == 200
    job = response.json()
    assert job["status"] == "succeeded"
    refreshed = client.get(f"/api/projects/{project['id']}").json()
    assert refreshed["audit_events"][0]["event_type"] == "rewrite.batch_completed"
    assert refreshed["audit_events"][0]["actor"] == "项目负责人"


def test_llm_prompt_keeps_late_chapter_coverage_for_long_novels():
    long_text = _make_long_novel(chapter_count=8, filler_repeats=520)
    prompt = _build_llm_prompt(long_text, "长篇测试", AdaptationStyle())

    assert "LONGFORM_CONTEXT" in prompt
    assert "第八章 终局密钥" in prompt
    assert "终局密钥藏在钟楼背面的裂缝里" in prompt
    assert "请按章节覆盖生成场景" in prompt
    assert len(prompt) < len(long_text)


def test_longform_pipeline_stage_is_reported_for_long_inputs():
    chapters = detect_chapters(_make_long_novel(chapter_count=8, filler_repeats=520))
    screenplay = build_fallback_screenplay(chapters, "长篇测试", AdaptationStyle())
    stages = {stage.id: stage for stage in screenplay.metadata.pipeline_stages}

    assert "stage_longform_context" in stages
    assert stages["stage_longform_context"].agent == "LongformChunker"
    assert stages["stage_longform_context"].output_count == 8


def _make_long_novel(chapter_count: int, filler_repeats: int) -> str:
    chapters: list[str] = []
    for index in range(1, chapter_count + 1):
        if index == chapter_count:
            title = "第八章 终局密钥"
            clue = "终局密钥藏在钟楼背面的裂缝里，只有林岚知道父亲留下的暗号。"
        else:
            title = f"第{index}章 漫长铺垫"
            clue = f"林岚在第{index}章继续追查旧案，线索暂时指向被封存的档案。"
        filler = "雨声压住走廊里的脚步，人物反复核对旧信和时间表。" * filler_repeats
        chapters.append(f"{title}\n\n{filler}\n\n{clue}")
    return "\n\n".join(chapters)


def _minimal_pdf_bytes() -> bytes:
    lines = ["第一章 雨夜来信", "第二章 旧街钟楼", "第三章 报社档案室"]
    commands = ["BT /F1 12 Tf 16 TL 72 720 Td"]
    for index, line in enumerate(lines):
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        commands.append(f"({escaped}) Tj")
        if index < len(lines) - 1:
            commands.append("T*")
    commands.append("ET")
    stream = " ".join(commands).encode("utf-8")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> /MediaBox [0 0 612 792] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    chunks = [b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"]
    offsets: list[int] = []
    current = len(chunks[0])
    for index, obj in enumerate(objects, start=1):
        offsets.append(current)
        chunk = f"{index} 0 obj\n".encode("ascii") + obj + b"\nendobj\n"
        chunks.append(chunk)
        current += len(chunk)
    xref_offset = current
    xref = [b"xref\n0 6\n0000000000 65535 f \n"]
    xref.extend(f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets)
    trailer = f"trailer\n<< /Root 1 0 R /Size 6 >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    return b"".join(chunks + xref + [trailer])
